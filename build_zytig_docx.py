# -*- coding: utf-8 -*-
"""
FCS-Zytig Sommer 2026 — Word-Generator.

Erzeugt die Vereinszeitung als .docx im Design-System des FC Schattdorf
(Vereinsrot #E63124, Ink #181818, Inter, Versal-Titel, rote Akzentbalken).

Aufbau gemäss Redaktionsvorgabe Sommer 26:
  Titelbild → Inhalt → Impressum →
  PRÄSIDENT (Vorwort, GV-Einladung) →
  SPONSORING (Platzhalter) →
  ADMINISTRATION (Platzhalter + Adressliste zur Prüfung) →
  EVENTS (Fasnacht, Jassturnier, Agenda) →
  SPORTCHEF (FCS 1–3, Senioren, Frauen 1+2, je mit Tabelle/Sponsoren) →
  JUNIOREN (Bericht, Fotoseiten, Sponsoren)

Aufruf:  python3 build_zytig_docx.py
Benötigt: python-docx, openpyxl, Pillow.
Liest:    Zytig Sommer 26/** (Inhalte), Statistik.xlsx (Tabellen),
          fotoklassifikation.txt (Foto-Auswahl, optional).
"""

import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageOps

# ---------------------------------------------------------------- Konstanten

REPO = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(REPO, "Zytig Sommer 26")
INSERATE = os.path.join(REPO, "assets", "img", "inserate")
OUT_DOCX = os.path.join(REPO, "FCS-Zytig Sommer 2026 ENTWURF.docx")
IMG_CACHE = os.path.join(REPO, "output", "img-cache")
FOTO_KLASSIFIKATION = os.environ.get(
    "FOTO_KLASSIFIKATION",
    os.path.join(REPO, "assets", "fotoklassifikation.txt"),
)

RED = RGBColor(0xE6, 0x31, 0x24)
RED_DEEP = RGBColor(0xD9, 0x26, 0x1C)
INK = RGBColor(0x18, 0x18, 0x18)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MUTED = RGBColor(0x52, 0x57, 0x5E)

HEX_RED = "E63124"
HEX_RED_DEEP = "D9261C"
HEX_INK = "181818"
HEX_MIST = "F6F2F1"
HEX_BORDER = "E5E2E1"

FONT = "Inter"
CONTENT_W = Cm(17.4)  # A4 minus 1.8 cm Rand links/rechts

# --------------------------------------------------------------- XML-Helfer


def _set_font(run, size=None, bold=None, color=None, caps=False, spacing=None,
              italic=None):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if caps:
        el = OxmlElement("w:caps")
        rpr.append(el)
    if spacing is not None:  # Laufweite in 1/20 pt
        el = OxmlElement("w:spacing")
        el.set(qn("w:val"), str(spacing))
        rpr.append(el)


def _shade_paragraph(par, hexcolor):
    ppr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    ppr.append(shd)


def _par_border(par, edge, hexcolor, size=24, space=4):
    ppr = par._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))  # 1/8 pt
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), hexcolor)
    pbdr.append(el)


def _shade_cell(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcpr.append(shd)


def _cell_borders(cell, edges, hexcolor=HEX_BORDER, size=4, style="single"):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), style)
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hexcolor)
        borders.append(el)


def _cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom),
                      ("start", left), ("end", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcpr.append(mar)


def _table_fixed(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tblpr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


def _no_space(par, before=0, after=0):
    par.paragraph_format.space_before = Pt(before)
    par.paragraph_format.space_after = Pt(after)


def _repeat_header_row(table):
    """Kopfzeile bei Seitenumbruch wiederholen."""
    trpr = table.rows[0]._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trpr.append(el)


def _update_fields_on_open(doc):
    """Word beim Öffnen auffordern, Felder (Inhaltsverzeichnis) zu aktualisieren."""
    settings = doc.settings.element
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


def _field(par, instr):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = " "
    run.append(text)
    fld.append(run)
    par._p.append(fld)


# ------------------------------------------------------------ Layout-Bausteine


def add_kicker(doc, text):
    """Rotes Kategorien-Label über einem Artikel (wie .kategorie im CSS)."""
    par = doc.add_paragraph()
    _no_space(par, before=10, after=2)
    run = par.add_run("  " + text + "  ")
    _set_font(run, size=8, bold=True, color=WHITE, caps=True, spacing=30)
    _shade_run(run, HEX_RED_DEEP)
    return par


def _shade_run(run, hexcolor):
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    rpr.append(shd)


def add_ressort_band(doc, ressort, verantwortlich=None):
    """Kapiteltrenner: roter Balken mit Ressortname, darunter Zuständigkeit.

    Wird als Heading 1 ausgezeichnet und erscheint damit im Inhaltsverzeichnis.
    """
    par = doc.add_paragraph()
    par.style = doc.styles["Heading 1"]
    _no_space(par, before=0, after=0)
    _shade_paragraph(par, HEX_RED)
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(6)
    run = par.add_run(" " + ressort)
    _set_font(run, size=21, bold=True, color=WHITE, caps=True, spacing=20)
    if verantwortlich:
        sub = doc.add_paragraph()
        _shade_paragraph(sub, HEX_INK)
        _no_space(sub, after=14)
        run = sub.add_run(" " + verantwortlich)
        _set_font(run, size=9, bold=True, color=WHITE, caps=True, spacing=25)
    else:
        spacer = doc.add_paragraph()
        _no_space(spacer, after=8)
    return par


def add_h2(doc, text, space_before=14):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(space_before)
    par.paragraph_format.space_after = Pt(6)
    par.paragraph_format.keep_with_next = True
    _par_border(par, "bottom", HEX_RED, size=20, space=3)
    run = par.add_run(text)
    _set_font(run, size=14.5, bold=True, color=INK, caps=True, spacing=10)
    return par


def add_h3(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(10)
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.keep_with_next = True
    run = par.add_run(text)
    _set_font(run, size=11, bold=True, color=RED_DEEP)
    return par


def add_body(doc, text, size=10, bold=False, color=INK, after=6, italic=False,
             line=1.15):
    par = doc.add_paragraph()
    _no_space(par, after=after)
    par.paragraph_format.line_spacing = line
    run = par.add_run(text)
    _set_font(run, size=size, bold=bold, color=color, italic=italic)
    return par


def add_signature(doc, text):
    par = add_body(doc, text, size=10, bold=True, after=10)
    return par


def add_placeholder_box(doc, titel, zeilen):
    """Grauer Platzhalterkasten mit rotem [PLATZHALTER]-Label."""
    table = doc.add_table(rows=1, cols=1)
    _table_fixed(table, [CONTENT_W])
    cell = table.rows[0].cells[0]
    _shade_cell(cell, HEX_MIST)
    _cell_borders(cell, ("top", "bottom", "start", "end"),
                  hexcolor=HEX_RED, size=8, style="dashed")
    _cell_margins(cell, top=140, bottom=140, left=180, right=180)
    par = cell.paragraphs[0]
    _no_space(par, after=4)
    run = par.add_run("[PLATZHALTER] ")
    _set_font(run, size=10, bold=True, color=RED_DEEP, caps=True, spacing=20)
    run = par.add_run(titel)
    _set_font(run, size=10, bold=True, color=INK)
    for zeile in zeilen:
        par = cell.add_paragraph()
        _no_space(par, after=2)
        run = par.add_run("•  " + zeile)
        _set_font(run, size=9.5, color=MUTED)
    doc.add_paragraph()
    return table


def add_sponsor_box(doc, eintraege, pruefen=True, logos=None):
    """Sponsorenhinweis unter einem Teambericht, optional mit Logo-Bildern."""
    table = doc.add_table(rows=1, cols=1)
    _table_fixed(table, [CONTENT_W])
    cell = table.rows[0].cells[0]
    _shade_cell(cell, HEX_MIST)
    _cell_borders(cell, ("start",), hexcolor=HEX_RED, size=24)
    _cell_margins(cell, top=90, bottom=90, left=180, right=180)
    par = cell.paragraphs[0]
    _no_space(par, after=2)
    run = par.add_run("Sponsoren")
    _set_font(run, size=8.5, bold=True, color=RED_DEEP, caps=True, spacing=25)
    if pruefen:
        run = par.add_run("   [PRÜFEN — Stand letzte Ausgabe]")
        _set_font(run, size=8, bold=True, color=MUTED)
    for label, namen in eintraege:
        par = cell.add_paragraph()
        _no_space(par, after=1)
        run = par.add_run(label + ": ")
        _set_font(run, size=9.5, bold=True, color=INK)
        run = par.add_run(namen)
        _set_font(run, size=9.5, color=INK)
    if logos:
        par = cell.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _no_space(par, before=6, after=2)
        for pfad, breite in logos:
            if os.path.exists(pfad):
                par.add_run().add_picture(pfad, width=Cm(breite))
                par.add_run("    ")
    doc.add_paragraph()
    return table


def prep_image(path, max_px=1100):
    """EXIF-Rotation anwenden, verkleinern, im Cache ablegen. → (pfad, b/h)."""
    os.makedirs(IMG_CACHE, exist_ok=True)
    key = re.sub(r"[^A-Za-z0-9]+", "_", os.path.basename(path)) + ".jpg"
    out = os.path.join(IMG_CACHE, key)
    img = ImageOps.exif_transpose(Image.open(path))
    ratio = img.width / img.height
    if not os.path.exists(out):
        img = img.convert("RGB")
        img.thumbnail((max_px, max_px))
        img.save(out, "JPEG", quality=78)
    return out, ratio


def add_photo_grid(doc, paths, cols=2, gap_after=True, max_h_cm=None):
    """Fotoraster: Bilder zeilenweise in unsichtbarer Tabelle.

    Hochformat-Bilder werden auf max_h_cm Höhe begrenzt (Seitenverhältnis
    bleibt erhalten), damit keine überhohen Zeilen entstehen.
    """
    paths = [p for p in paths if os.path.exists(p)]
    if not paths:
        return
    if max_h_cm is None:
        max_h_cm = {1: 12.0, 2: 8.5, 3: 7.0}.get(cols, 8.0)
    width_cm = CONTENT_W.cm - 0.4 * (cols - 1)
    cell_w = Cm(width_cm / cols)
    rows = (len(paths) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    _table_fixed(table, [cell_w] * cols)
    for idx, path in enumerate(paths):
        cell = table.rows[idx // cols].cells[idx % cols]
        _cell_margins(cell, top=40, bottom=40, left=40, right=40)
        par = cell.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _no_space(par)
        cached, ratio = prep_image(path)
        w_cm = cell_w.cm - 0.25
        if w_cm / ratio > max_h_cm:  # zu hoch → über Höhe skalieren
            par.add_run().add_picture(cached, height=Cm(max_h_cm))
        else:
            par.add_run().add_picture(cached, width=Cm(w_cm))
    if gap_after:
        doc.add_paragraph()


def add_centered_image(doc, path, width_cm=None, height_cm=None, before=6,
                       after=6, break_before=False):
    """Einzelbild zentriert einfügen (Inserate, Logos, Flyer)."""
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(before)
    par.paragraph_format.space_after = Pt(after)
    if break_before:
        par.paragraph_format.page_break_before = True
    kwargs = {}
    if width_cm:
        kwargs["width"] = Cm(width_cm)
    if height_cm:
        kwargs["height"] = Cm(height_cm)
    par.add_run().add_picture(path, **kwargs)
    return par


def add_inserat_hinweis(doc):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(par, before=2, after=2)
    run = par.add_run("[PRÜFEN] Inserat aus letzter Ausgabe übernommen — "
                      "Aktualität durch Sponsoringverantwortlichen bestätigen.")
    _set_font(run, size=7.5, color=MUTED, italic=True)
    return par


def extrahiere_fasnacht_bilder():
    """Eingebettete Bilder des Fasnacht-Berichts, gruppiert nach Abschnitt."""
    import docx as _docx
    pfad = os.path.join(SRC, "Events",
                        "BERICHT_FCS Zeitung Fasnachtsanlässe 2026.docx")
    ziel = os.path.join(IMG_CACHE, "fasnacht")
    os.makedirs(ziel, exist_ok=True)
    d = _docx.Document(pfad)
    ns = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    rel = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
    gruppen = {"vereinsfasnacht": [], "sbu": [], "jungalt": [], "dank": []}
    seq = 0
    for i, par in enumerate(d.paragraphs):
        for blip in par._p.findall(f".//{ns}blip"):
            part = d.part.related_parts[blip.get(f"{rel}embed")]
            seq += 1
            fn = os.path.join(ziel, f"fasnacht_{seq:02d}.png")
            if not os.path.exists(fn):
                with open(fn, "wb") as fh:
                    fh.write(part.blob)
            if i <= 20:
                gruppen["vereinsfasnacht"].append(fn)
            elif i <= 25:
                gruppen["sbu"].append(fn)
            elif i <= 34:
                gruppen["jungalt"].append(fn)
            else:
                gruppen["dank"].append(fn)
    return gruppen


def add_liga_tabelle(doc, titel, zeilen, highlight):
    """Ligatabelle: Rang, Verein, Sp, S, U, N, SP, Tore, Pkt."""
    add_h3(doc, titel)
    header = ["Rg", "Verein", "Sp", "S", "U", "N", "Str", "Tore", "Pkt"]
    widths = [Cm(1.1), Cm(6.5), Cm(1.1), Cm(1.1), Cm(1.1), Cm(1.1),
              Cm(1.3), Cm(2.4), Cm(1.7)]
    table = doc.add_table(rows=1 + len(zeilen), cols=len(header))
    _table_fixed(table, widths)
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        _shade_cell(cell, HEX_RED)
        _cell_margins(cell)
        par = cell.paragraphs[0]
        _no_space(par)
        if i >= 2:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(text)
        _set_font(run, size=8.5, bold=True, color=WHITE, caps=True)
    for r, zeile in enumerate(zeilen, start=1):
        is_hl = highlight.lower() in zeile[1].lower()
        for i, val in enumerate(zeile):
            cell = table.rows[r].cells[i]
            _cell_margins(cell)
            _cell_borders(cell, ("bottom",))
            if is_hl:
                _shade_cell(cell, HEX_MIST)
            par = cell.paragraphs[0]
            _no_space(par)
            if i >= 2 or i == 0:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par.add_run(str(val))
            _set_font(run, size=9, bold=is_hl, color=INK)
    _repeat_header_row(table)
    par = doc.add_paragraph()
    _no_space(par, before=2, after=10)
    run = par.add_run("Sp Spiele · S Siege · U Unentschieden · N Niederlagen · "
                      "Str Strafpunkte · Pkt Punkte")
    _set_font(run, size=7.5, color=MUTED)


def add_stat_tabelle(doc, titel, header, zeilen, widths, legende=None):
    """Generische Statistik-Tabelle (Einsatzstatistik etc.)."""
    add_h3(doc, titel)
    table = doc.add_table(rows=1 + len(zeilen), cols=len(header))
    _table_fixed(table, widths)
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        _shade_cell(cell, HEX_INK)
        _cell_margins(cell)
        par = cell.paragraphs[0]
        _no_space(par)
        if i >= 1:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(text)
        _set_font(run, size=8, bold=True, color=WHITE, caps=True)
    for r, zeile in enumerate(zeilen, start=1):
        for i, val in enumerate(zeile):
            cell = table.rows[r].cells[i]
            _cell_margins(cell, top=20, bottom=20)
            _cell_borders(cell, ("bottom",))
            par = cell.paragraphs[0]
            _no_space(par)
            if i >= 1:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par.add_run(str(val))
            _set_font(run, size=8.5, color=INK)
    _repeat_header_row(table)
    if legende:
        par = doc.add_paragraph()
        _no_space(par, before=2, after=10)
        run = par.add_run(legende)
        _set_font(run, size=7.5, color=MUTED)
    else:
        doc.add_paragraph()


def add_agenda(doc, eintraege):
    for datum, titel, ort in eintraege:
        par = doc.add_paragraph()
        _no_space(par, after=0)
        par.paragraph_format.space_before = Pt(4)
        _par_border(par, "bottom", HEX_BORDER, size=4, space=3)
        run = par.add_run(datum.ljust(22))
        _set_font(run, size=10, bold=True, color=RED_DEEP)
        run = par.add_run(titel)
        _set_font(run, size=10, bold=True, color=INK)
        if ort:
            run = par.add_run("  ·  " + ort)
            _set_font(run, size=9.5, color=MUTED)


# ------------------------------------------------------------------- Kopf/Fuss


def build_masthead(section):
    header = section.header
    header.is_linked_to_previous = False
    for par in list(header.paragraphs):
        par._p.getparent().remove(par._p)
    table = header.add_table(rows=1, cols=1, width=CONTENT_W)
    _table_fixed(table, [CONTENT_W])
    cell = table.rows[0].cells[0]
    _shade_cell(cell, HEX_INK)
    _cell_margins(cell, top=80, bottom=80, left=160, right=160)
    par = cell.paragraphs[0]
    _no_space(par)
    par.paragraph_format.tab_stops.add_tab_stop(
        Cm(16.7), WD_TAB_ALIGNMENT.RIGHT)
    run = par.add_run("FCS-ZYTIG")
    _set_font(run, size=13, bold=True, color=WHITE, spacing=30)
    run = par.add_run("\tSOMMER 2026 · FC SCHATTDORF")
    _set_font(run, size=8, bold=True, color=WHITE, spacing=25)
    line = header.add_paragraph()
    _shade_paragraph(line, HEX_RED)
    _no_space(line)
    line.paragraph_format.line_spacing = Pt(3)

    footer = section.footer
    footer.is_linked_to_previous = False
    for par in list(footer.paragraphs):
        par._p.getparent().remove(par._p)
    par = footer.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _par_border(par, "top", HEX_RED, size=12, space=4)
    run = par.add_run("FCS-Zytig · Sommer 2026 · Seite ")
    _set_font(run, size=8, bold=True, color=MUTED, caps=True, spacing=20)
    _field(par, "PAGE")
    for r in par.runs:
        _set_font(r, size=8, bold=True, color=MUTED)


# --------------------------------------------------------------------- Inhalte

VORWORT = [
    "Liebe Mitglieder, Freunde und Unterstützer des FC Schattdorf",
    "Eine weitere Saison liegt hinter uns – eine Saison mit Höhen und Tiefen, aber vor allem mit vielen Menschen, die sich mit grossem Engagement für unseren FC Schattdorf eingesetzt haben. Ein Verein wie unserer lebt nicht nur von den Resultaten auf dem Platz, sondern vor allem von den vielen Personen, die im Hintergrund Verantwortung übernehmen und ihre Zeit für Rot-Schwarz einsetzen. Mein herzlicher Dank gilt deshalb meinen Vorstandskolleginnen und -kollegen sowie allen Trainerinnen und Trainern, Betreuern, Funktionären, Schiedsrichtern und den zahlreichen weiteren ehrenamtlich engagierten Helferinnen und Helfern. Ohne euren Einsatz wäre ein aktives Vereinsleben in dieser Form schlicht nicht möglich. Danke, dass ihr unseren Verein mittragt und weiterentwickelt.",
    "Auch sportlich dürfen wir auf eine insgesamt erfreuliche Saison zurückblicken. Unsere 1. Mannschaft hat nach einem schwierigen Start in die Rückrunde und der daraus resultierenden Trainerentlassung Charakter bewiesen. Das Team hat sich gefangen und die Meisterschaft versöhnlich abgeschlossen. Auch unsere 2. Mannschaft durfte lange vom Aufstieg träumen. Zwar hat es am Ende nicht ganz gereicht, doch die Entwicklung stimmt zuversichtlich und die Formkurve zeigt klar in die richtige Richtung. Ebenso erfreulich präsentieren sich unsere Junioren- sowie die weiteren Aktivteams. Über alle Alters- und Leistungsstufen hinweg wird engagiert gearbeitet, und wir dürfen mit den gezeigten Leistungen zufrieden sein.",
    "Ein Highlight war in diesem Jahr die erstmalige Durchführung unseres Kick-in-one-Fussballgolfs. Dass ein neuer Anlass nicht vom ersten Moment an perfekt funktioniert, liegt in der Natur der Sache. Wir sind uns bewusst, dass insbesondere das Timing beim Sponsorensammeln Verbesserungspotenzial hat, sodass weniger Überschneidungen mit dem Grümpi-Losverkauf entstehen. Die entsprechenden Rückmeldungen nehmen wir ernst und werden sie in unsere zukünftigen Planungen einfliessen lassen. Gleichzeitig hat uns das sehr positive Echo der Teilnehmenden bestätigt, dass sich der Aufwand gelohnt hat. Für viele war das Kick-in-one-Fussballgolf ein tolles Erlebnis, weshalb wir überzeugt sind, diesen Anlass auch künftig wieder durchführen zu wollen. Ein herzliches Dankeschön gilt an dieser Stelle meinen Vorstandskollegen sowie Patrik Müller für die grosse Unterstützung bei der Organisation und Durchführung.",
    "Vielleicht stellt sich der eine oder die andere die Frage, weshalb wir als Verein solche zusätzlichen Anlässe überhaupt durchführen. Die Antwort ist einfach: Sie ermöglichen uns Investitionen, die unseren Mitgliedern langfristig zugutekommen. Auch im kommenden Vereinsjahr möchten wir verschiedene Projekte realisieren. Geplant sind unter anderem ein neuer Veloparkplatz mit Platz für rund 78 Fahrräder und 25 Kickboards beim Eingangsbereich, die Absperrung der Zufahrt beim Galgenwäldli, damit diese als Blaulichtzufahrt jederzeit frei bleibt, der Einbau einer Enthärtungsanlage zur Verlängerung der Lebensdauer unseres Boilers sowie zusätzliche Lagermöglichkeiten auf der Grundmatte – voraussichtlich mittels eines Containers – und damit verbunden auch der Erwerb von neuen, hochwertigen Fussballtoren für unsere Nachwuchsabteilung. Solche Investitionen kommen unserem gesamten Verein zugute und sind nur dank zusätzlicher Einnahmen sowie der Unterstützung unserer Mitglieder, Sponsoren und Gönner realisierbar.",
    "Ein weiterer Schwerpunkt wird in der kommenden Saison die Teilnahme am Programm «SFV Quality Club» sein. Dieses Programm unterstützt Vereine dabei, ihre Strukturen und Prozesse gezielt weiterzuentwickeln und sich in Bereichen wie Vereinsführung, Organisation, Ehrenamt, Nachwuchs oder Kommunikation nachhaltig zu verbessern. Genau das entspricht auch unserem Selbstverständnis: Wir wollen uns nicht auf Erreichtem ausruhen, sondern unseren Verein kontinuierlich weiterentwickeln. Ich freue mich darauf, diesen Weg gemeinsam mit dem Vorstand und vielen engagierten Vereinsmitgliedern zu gehen und dabei neue Impulse für die Zukunft des FC Schattdorf zu gewinnen.",
    "Ich danke euch allen herzlich für euren Einsatz und eure Unterstützung für unseren FC Schattdorf.",
]

TRAKTANDEN = [
    ("1.", "Begrüssung", []),
    ("2.", "Anwesenheitskontrolle", []),
    ("3.", "Wahl der Stimmenzähler", []),
    ("4.", "Protokoll der 92. Generalversammlung vom 22. August 2025", []),
    ("5.", "Mutationen", []),
    ("6.", "Ressortberichte", ["6.1  Präsident", "6.2  Sportchef",
                               "6.3  Juniorenobmann"]),
    ("7.", "Finanzen", ["7.1  Kassabericht per 30. Juni 2026",
                        "7.2  Bericht der Revisoren",
                        "7.3  Genehmigung der Jahresrechnung 2025/26"]),
    ("8.", "Wahlen", ["8.1  Präsident", "8.2  Sportchef", "8.3  Spiko",
                      "8.4  Administration", "8.5  Marketing", "8.6  Revisor"]),
    ("9.", "Festsetzung der Mitgliederbeiträge 2026/27", []),
    ("10.", "Vorstellung und Beschlussfassung Budget 2026/27", []),
    ("11.", "Anträge", ["11.1  Seitens des Vorstands",
                        "11.2  Seitens der Mitglieder"]),
    ("12.", "Ehrungen", []),
    ("13.", "Verschiedenes", []),
    ("14.", "Nachtessen", []),
]

FCS1 = [
    "Nach der punktemässig sehr erfolgreichen Vorrunde mit 9 Siegen aus 13 Spielen und insgesamt 28 Punkten und dazu der Serie von sechs Siegen zum Schluss ging man mit viel Selbstvertrauen in die Rückrunde.",
    "Es war wiederum das Ziel, an die erfolgreiche Vorrunde anzuknüpfen, um möglichst lange an der Tabellenspitze dabei sein zu können. Nach einer erfolgreichen Vorbereitungsphase, unter anderem mit Trainingsspielen gegen spielstarke Teams aus der 2. Liga interregional, startete man nicht wie erhofft in die Rückrunde. Sämtliche 4 Startspiele verlor unser Fanionteam, und auch die beiden darauffolgenden Spiele gegen Perlen und Sarnen vermochte man nicht zu gewinnen und erreichte jeweils nur ein Remis. Als man schliesslich auch im 7. Spiel in Folge gegen Brunnen keinen Sieg einholte und wiederum als Verlierer vom Platz gehen musste, entschieden wir uns, zusammen mit Trainer Emanuele de Masi, dass es in dieser Art und Weise nicht weitergehen kann und wir auf dem Trainerposten eine neue Lösung suchen müssen.",
    "Ad interim übernahm Thomas Zberg kurzerhand wiederum das Amt des Trainers zusammen mit Coach Reto Infanger, und der Erfolg kehrte bald zurück zum FC Schattdorf. Leider verlor man noch das erste Spiel unglücklich in der Nachspielzeit auswärts bei den Kickers. Doch danach konnte man in allen Spielen bis zum Saisonende punkten und gewann sogar alle 4 letzten Spiele.",
    "Trotz all den Schwierigkeiten darf man wiederum von einer guten Saison sprechen. Aus 26 Spielen erzielte man 43 Punkte und beendete die Saison auf dem 5. Tabellenrang. An dieser Stelle gehört sicherlich ein Kompliment an das Team. Die Reaktion, welche das Team nach dem nicht einfachen Rückrundenstart gezeigt hat, verdient Respekt und zeugt von einem tollen Mannschaftsgefüge.",
    "Aus Vereinssicht war die Saison sicherlich auch sehr anspruchsvoll. Es ist nie gewünscht oder gesucht, während der Saison eine neue Lösung auf Trainerebene finden zu müssen. Doch das Wohl und die DNA des Vereins stehen immer im Vordergrund, und dieses Ziel wurde und muss immer oberste Priorität haben. Der FC Schattdorf steht dafür ein, seinen Spielern ein Ambiente zu bieten, um Spass zu haben und gleichermassen erfolgreich zu sein. Er steht natürlich auch dafür ein, junge Spieler zu fördern und zu fordern und ihnen die Möglichkeit zu geben, sich in der 1. Mannschaft zu integrieren. Und ganz wichtig: Wir stehen dafür ein, unser intaktes Vereinsleben stets aufrechtzuerhalten. Die Zuschauer und all die vielen Helfer sollen sich mit der 1. Mannschaft identifizieren können. Nur das ist der richtige Weg und führt uns auch in Zukunft in erfolgreiche Zeiten.",
    "An dieser Stelle möchte ich mich noch kurz bedanken bei:",
    "•  Thomas Zberg für die sofortige Zusage als Trainer während der nicht einfachen Situation;\n•  Reto Infanger für die tolle Zusammenarbeit und Unterstützung während der letzten 5 Saisons;\n•  Sandro, Boeri, Skander, Gian und Livio für die Vereinstreue über die letzten Jahre;\n•  Mathias Lussmann für seinen grossen Einsatz in den letzten 3 Saisons als Trainer der 2. Mannschaft;\n•  der Sportkommission für die Unterstützung;\n•  … und ganz wichtig: allen Zuschauern, welche unsere Teams immer wunderbar unterstützen, sei es im Stadion Grüner Wald oder auch auswärts – die Teams schätzen das sehr!",
    "Ab der Saison 26/27 übernimmt Saverio La Bella die Funktion als Cheftrainer der 1. Mannschaft, zusammen mit Thomas Zberg, welcher dann die Funktion des Assistenztrainers übernehmen wird.",
]

FCS2 = [
    "Die 2. Mannschaft des FC Schattdorf darf auf eine starke und insgesamt sehr erfolgreiche Saison 2025/2026 zurückblicken. Bereits in der Vorbereitung zeigte das Team gute Ansätze und tankte mit überzeugenden Testspielsiegen gegen Hergiswil II, Brunnen III und Erstfeld II viel Selbstvertrauen. Besonders die offensive Durchschlagskraft und die mannschaftliche Geschlossenheit waren dabei deutlich sichtbar.",
    "Der Start in die Rückrunde verlief nahezu optimal. Mit einem klaren 3:0-Derbysieg gegen Altdorf II und einem überzeugenden 4:1-Heimerfolg gegen Ibach II setzte die Mannschaft früh ein Ausrufezeichen und bestätigte ihre Ambitionen für die zweite Saisonhälfte. Danach folgte jedoch mit dem 1:1 auswärts gegen das spätere Schlusslicht FC Sins b ein erster Dämpfer, bei dem wichtige Punkte liegen gelassen wurden.",
    "Im weiteren Verlauf wechselten sich starke Leistungen mit einzelnen Rückschlägen ab. Besonders die Heimniederlagen gegen Aegeri II, Hünenberg I und Brunnen II verhinderten eine noch bessere Klassierung. Trotzdem bewies die Mannschaft immer wieder Charakter und reagierte auf schwierige Phasen mit wichtigen Erfolgen. Die Auswärtssiege gegen Dietwil, Weggis und Küssnacht II unterstrichen die Qualität und Mentalität des Teams. Vor allem der 1:0-Erfolg beim späteren Meister Weggiser SC darf sicherlich als eines der Highlights der Saison bezeichnet werden.",
    "Bis zum zweitletzten Spieltag durfte die Mannschaft sogar auf die Teilnahme an den Aufstiegsspielen hoffen. Schlussendlich beendete der FC Schattdorf die Saison mit starken 35 Punkten auf dem hervorragenden 4. Tabellenrang – punktgleich mit dem Drittplatzierten FC Ibach II. Damit bestätigte das Team eindrucksvoll seine Zugehörigkeit zur Spitzengruppe der Liga.",
    "Neben den sportlichen Fortschritten bleiben vor allem auch die gemeinsamen Emotionen, die intensiven Spiele und der besondere Teamgeist in Erinnerung. Die Mannschaft hat in dieser Saison gezeigt, welches Potenzial in ihr steckt und dass sie die Qualität besitzt, künftig auch den nächsten Schritt machen zu können.",
    "Zum Abschluss bedankt sich die 2. Mannschaft herzlich bei allen Zuschauerinnen und Zuschauern, Helferinnen und Helfern sowie dem gesamten Umfeld des FC Schattdorf für die Unterstützung während der gesamten Saison.",
]

FCS3 = [
    "Die Saison der dritten Mannschaft des FC Schattdorf in der 5. Liga glich in den vergangenen Monaten einer wilden Achterbahnfahrt. Nach einer von Rückschlägen geprägten Rückrunde verabschiedete sich das Team mit einem sensationellen Paukenschlag in die Sommerpause. Am Ende steht der achte Tabellenplatz – und eine Fülle an Geschichten, die so nur der Amateurfussball schreiben kann.",
    ("H3", "Erst Tag der offenen Tür, dann Kabinen-Chaos"),
    "Der Start in die Frühlingsrunde verlief für die Schattdorfer alles andere als nach Wunsch. Besonders die herben Klatschen gegen den FC Flüelen (1:6) und den späteren Gruppen-Dominator FC Baar II (0:8) zeigten der Defensive temporär gnadenlos die Grenzen auf. Als dann auch noch die Partien gegen Steinhausen III, Erstfeld II und Muotathal verloren gingen, drohte das Team gänzlich im Tabellenkeller zu versinken.",
    "Doch die Mannschaft bewies Charakter. Am 31. Mai erlöste man sich im Kellerduell gegen den FC Ibach IV mit einem hart umkämpften 2:1-Heimsieg. Dieser Dreier gab spürbaren Auftrieb für den grossen Showdown am letzten Spieltag: Am Sonntagmorgen empfing man den Tabellenzweiten FC Brunnen III. In einer hochdramatischen Partie wuchsen die Schattdorfer über sich hinaus und bodigten den Favoriten mit 3:2. Man munkelt heute noch in den Katakomben des Sportplatzes, dass bei diesem Coup nicht alles mit ganz rechten Dingen zuging.",
    ("H3", "Wer hat noch nicht, wer will nochmal?"),
    "Für die grösste Kuriosität der Saison sorgte zweifellos die Position zwischen den Pfosten. Trainer Yannic Jäger musste bei der Torhüterfrage seine gesamte Kreativität aufbieten, da das Trikot mit der Nummer 1 regelrecht zum Wanderpokal mutierte: Insgesamt sechs verschiedene Akteure hüteten im Saisonverlauf das Tor. Neben Ex-Stammkeeper Yannic Jäger streiften sich auch Ranadan Fragnito, Kay Schillig und Matteo Zberg die Handschuhe über. In der Rückrunde kam sogar noch Claudio Pfyl hinzu, der den Kasten in insgesamt 7 Einsätzen hütete. Das absolute Highlight folgte jedoch im letzten Spiel gegen Brunnen: Da kein erstklassiger Goalie zur Verfügung stand, ging Stürmer Mariano Prandi «gezwungenermassen» ins Tor – und hielt den sensationellen Sieg mit Bravour fest.",
    ("H3", "Die alte Garde"),
    "In der teaminternen Statistik ragt vor allem ein Name heraus. In der Defensive erwies sich Reto Bissig als der unumstrittene Dauerbrenner: «Papacito Reti» stand in 17 von 18 ausgewerteten Partien in der Startelf und steuerte zudem zwei Treffer bei. Besondere Anerkennung gebührt dem Edeljoker der Mannschaft, Patrick Gamma. Er kam stolze 7-mal als Einwechselspieler von der Bank und bewies eindrücklich, dass Klasse zeitlos ist: Der Routinier überzeugte mit seiner enormen Erfahrung sowie Fitness, sorgte jedes Mal sofort für frischen Wind in der Partie und krönte seine starken 9 Auftritte im Laufe der Saison mit insgesamt vier wichtigen Treffern. Auf ihn konnten die Schattdorfer immer zählen, wenn Not am Mann war.",
    ("H3", "Ein versöhnliches Fazit und bittere Abschiede"),
    "Mit 10 Punkten und einem Torverhältnis von 24:69 Toren schliesst der FC Schattdorf 3 die Spielzeit auf dem 8. Rang ab und liess den FC Ibach IV hinter sich. Was dem Team in manchen Phasen an defensiver Stabilität fehlte, machte es durch vorbildliche Fairness wett: Mit einer verschwindend geringen Anzahl an gelben Karten gehörte man zu den fairsten Teams der Liga.",
    "Nach dem heroischen Erfolg gegen den Tabellenzweiten Brunnen verabschiedet sich die Mannschaft in die verdiente Sommerpause. Wenn der Schwung aus dem Saisonfinale konserviert werden kann, ist mit den Schattdorfern in der nächsten Spielzeit definitiv weiter oben zu rechnen.",
    "Zum Saisonende muss die Mannschaft leider auch einige schmerzliche Abgänge verkraften. Während ein Teil der Routiniers dem unaufhaltsamen Ruf der Biologie folgt und die Fussballschuhe künftig eine Altersklasse höher bei den Senioren schnürt, zieht es andere Akteure in den vorgezogenen, fussballerischen Vorruhestand. Böse Zungen behaupten, der verlockende Ruf der heimischen Couch am Sonntagvormittag war am Ende einfach lauter als die Motivationsansprachen des Trainers. Sie alle werden auf und neben dem Platz schmerzlich vermisst werden.",
]

SENIOREN = [
    "Die Rückrunde begann ähnlich verheissungsvoll wie ein Wecker am Montagmorgen: Man weiss, was kommt – und freut sich trotzdem nicht darauf. Die Senioren Uri mussten in den ersten vier Spielen gleich gegen die Top-4 ran. Ein Auftaktprogramm, das an jene Cupauslosungen erinnert, bei denen man hofft: «Bitte nicht die Grossen …» – und dann natürlich genau die Grossen bekommt. Ergebnis: vier Niederlagen am Stück, wie schon in der Vorrunde. Konstanz ist eben auch eine Qualität.",
    "Im ersten Spiel gegen Buochs hielt man in der ersten Halbzeit richtig gut mit und ging sogar in Führung. Es fühlte sich an wie diese magischen Momente, wenn ein Aussenverteidiger plötzlich einen Übersteiger auspackt, den er seit den C-Junioren nie mehr probiert hat – und er gelingt. Doch dann kamen die Eigenfehler. Am Ende stand ein 1:4.",
    "In Cham gab es dann gar nichts zu holen. Das Spiel erinnerte an jene Tage, an denen man schon beim Einlaufen merkt: «Heute wird’s zäh.» 0:4 – und man war froh, dass es nicht noch schlimmer wurde.",
    "Im dritten Spiel gegen den späteren Aufsteiger SC Goldau zeigte Uri lange eine starke Leistung. Es war ein Hin und Her, wie wenn zwei Teams sich gegenseitig den Ball nicht gönnen wollen. Doch vorne liess man zu viel liegen – und wie jeder alte Fussballer weiss: «Wenn du sie nicht machst, macht sie der andere.» Genau so kam es. 2:4.",
    "In Hünenberg verschlief man die ersten 25 Minuten komplett. Es war, als hätte die Mannschaft kollektiv den Anpfiff mit der Platzbewässerung verwechselt. Zwei unnötige Penaltys später stand man 0:2 hinten. In der zweiten Halbzeit kämpfte sich die Bär-Elf zurück und glich zum 2:2 aus – nur um dann mit dem Schlusspfiff das 2:3 zu kassieren. Ein klassischer «Wir waren eigentlich schon in der Kabine»-Moment.",
    "Erst im fünften Spiel der Rückrunde platzte der Knoten – und wie. Ein 4:0 gegen den SC Kriens, ein Sieg, der sich anfühlte wie der erste warme Frühlingstag nach einem langen Winter. Plötzlich lief der Ball, plötzlich lief alles.",
    "Doch im nächsten Spiel gegen Perlen regnete es so stark, dass man fast Schwimmflossen gebraucht hätte. Der Ball blieb stehen, die Laune auch – 1:2-Niederlage. Ein Spiel, das man eher in die Kategorie «Wetterlotterie» einordnet.",
    "Beim Familien-Spiel in Altdorf, bei Temperaturen, die selbst Kunstrasen zum Schmelzen bringen könnten, gelang dann der nächste Dreier. 2:1 gegen Rotkreuz. Die Hitze war so brutal, dass selbst der Schiedsrichter aussah, als würde er gleich um eine Trinkpause bitten.",
    "Im zweitletzten Spiel gegen Weggis zeigte die Urner-Elf dann Fussball, der so flüssig war wie ein frisch geölter Rasenmäher. 4:1 – ein Spiel, bei dem man merkte: Heute passt einfach alles. Selbst die Einwürfe sahen elegant aus.",
    "Die Rückrunde endete dann, wie sie begonnen hatte – mit einer Niederlage. 2:5 in Alpnach. Ein Spiel, das man wohl unter «Lehrblätz» verbucht. Oder wie ein erfahrener Seniorenspieler mal sagte: «Manchmal gewinnt man, manchmal lernt man. Und manchmal lernt man halt ein bisschen mehr.»",
    "Zum Schluss ein grosses Dankeschön an den Trainerstaff und die Clubhausteams der drei Urner Vereine für ihren unermüdlichen Einsatz über die gesamte Saison hinweg. Ohne euch wäre das Ganze nur halb so organisiert, halb so motiviert – und vermutlich doppelt so chaotisch.",
]

FRAUEN1 = [
    "Die Saison 2025/26 stellte das Team Uri Frauen 1 vor besondere Herausforderungen, denn die Rückrunde musste ohne offizielle/n Trainer/in bestritten werden. Die Verantwortung übernahm der Spielerinnenrat, bestehend aus den fünf Aktivspielerinnen Dominique Scheiber, Anita Arnold, Svenja Arnold, Julia Novacic und Noreen Häfliger. Gemeinsam organisierten sie die Vorbereitung, Trainings und Spiele. Das Team Uri 1 war somit weitgehend auf sich allein gestellt – eine Situation, die grossen Zusammenhalt und Eigeninitiative erforderte.",
    "Die Vorbereitung im Winter wurde intensiv gestaltet und durch drei Trainingsspiele ergänzt. Aufgrund von wetterbedingten Spielverschiebungen startete das Team Uri 1 nicht wie üblich in die Meisterschaft, sondern direkt mit dem Cup-Viertelfinal gegen Willisau. Dort zeigte man ein überzeugendes Spiel über 90 Minuten und gewann deutlich mit 5:0.",
    "In den darauffolgenden vier Meisterschaftsspielen konnte jedoch nicht ganz an diese Leistung angeknüpft werden. Die Bilanz: zwei Siege, ein Unentschieden und eine Niederlage.",
    "Ein weiteres Highlight folgte im Cup-Halbfinal gegen Obwalden. Mit einem souveränen 7:2-Sieg zeigte das Team Uri 1 wohl seine beste Saisonleistung und zog verdient in den Cupfinal ein. Die Mannschaft dominierte das Spiel über weite Strecken und überzeugte sowohl spielerisch als auch kämpferisch.",
    "Anschliessend folgten drei weitere Meisterschaftsspiele, von denen zwei gewonnen werden konnten, darunter ein wichtiger Sieg gegen die Leaderinnen aus Balerna.",
    "Der Höhepunkt der Saison war klar der Cupfinal gegen Frauenfussball Seetal. Uri nahm sich vor, erneut die starken Cup-Leistungen abzurufen. Der Start gelang nach Mass: Bereits nach fünf Minuten ging das Team Uri 1 mit 1:0 in Führung. In einer grösstenteils ausgeglichenen Partie erspielte man sich weitere Chancen, konnte diese jedoch nicht nutzen. Seetal hingegen war bei Standardsituationen effizient und erzielte drei Tore, wodurch das Spiel leider mit 3:1 verloren ging.",
    "Zum Saisonabschluss standen danach noch zwei Meisterschaftsspiele auf dem Programm, mit einem Sieg und einer Niederlage.",
    ("H3", "Fazit der Saison"),
    "Trotz schwieriger Umstände ohne Trainer/in zeigten die Frauen vom Team Uri 1 grossen Einsatz und Zusammenhalt. Besonders im Cup konnten die besten Leistungen abgerufen werden, auch wenn die Krönung im Final leider ausblieb. In der Meisterschaft fehlten oft nur Kleinigkeiten zum Erfolg. Am Ende schliesst das Team Uri Frauen 1 die Saison auf dem 5. Tabellenplatz ab.",
    "Zum Saisonende werden mit Dominique Scheiber und Monika Mulle zwei langjährige Spielerinnen verabschiedet. Vielen Dank für euren grossartigen Einsatz – ihr werdet der Mannschaft fehlen!",
    "Ein besonderer Dank gilt auch dem Spielerinnenrat. Ohne seine Initiative wäre die Saison in dieser Form nicht möglich gewesen. Ebenfalls ein grosses Dankeschön an alle Unterstützer/innen, die das Team im Training und bei Spielen begleitet haben: Fabrice, Jasmin, Tim, Ramon, Igor & Sascha, Fabio, Ronny & Augusto.",
    "Die kommende Saison bringt strukturelle Veränderungen mit sich: Die Frauen stellen neu nur noch eine Aktivmannschaft, welche von Fabrice Arnold als neuem Trainer geführt wird. Die Mannschaft freut sich über den frischen Impuls von Fabrice und blickt gespannt und motiviert auf die kommende Saison.",
]

FRAUEN2 = [
    "Die zweite Mannschaft des Team Uri hat die Vorrunde der Meisterschaft mit viel Einsatz und Kampfgeist bestritten. Die Mannschaft startete die Saison mit einem neuen Trainerteam und einigen neuen Gesichtern. Anfangs brachte das eine grosse Umstellung mit sich. Den Urnerinnen gelang es jedoch, sich schnell einzuspielen und so bereit für die Meisterschaft zu sein.",
    "Im Verlauf der Vorrunde zeigten sich die Urnerinnen immer wieder von ihrer kämpferischen Seite. Es wurde versucht, Chancen herauszuspielen und das Spiel zu kontrollieren. Dabei zeigte sich, dass oftmals das letzte Quäntchen Glück fehlte – sei es vor dem gegnerischen Tor oder in der eigenen Verteidigung. So entstanden, trotz guter Aktionen und hoher Einsatzbereitschaft, nicht immer die erhofften Ergebnisse. Durch diese Niederlagen und Rückschläge liessen sich die Urnerinnen jedoch nicht unterkriegen. Besonders auffallend war, dass die Spielerinnen gemeinsam kämpften und sich gegenseitig halfen. Auch in hektischen Momenten versuchten sie zusammenzuspielen und glaubten stetig an den Sieg.",
    "Die Resultate waren über die Saison verteilt ein wenig schwankend. Trotzdem gelangen ihnen 2 Siege und 1 Unentschieden. Schlussendlich belegte die 2. Mannschaft vom Team Uri den 5. Rang in der Tabelle. Dabei konnten sich die Urnerinnen im Verlauf der Vorrunde positiv entwickeln, wuchsen als Team zusammen und zeigten eine starke Willenskraft. In der Winterpause wurde an Technik und Kondition gearbeitet.",
    "Die Rückrunde der Saison 2025/26 war für das Team Uri Frauen II von Höhen und Tiefen geprägt. Trotz einiger Herausforderungen zeigte die junge und engagierte Mannschaft während der gesamten Saison grossen Einsatz, Teamgeist und ihr vorhandenes Potenzial.",
    "Die Rückrunde begann mit zwei Unentschieden gegen den FC Aegeri und den späteren Gruppensieger Zug 94. Zu den Höhepunkten gehörten der überzeugende 6:2-Heimsieg gegen den FC Malcantone sowie der 3:2-Auswärtserfolg gegen den FC Baar III. Nach der 0:2-Niederlage gegen den SC Schwyz II setzte das Team mit einem 2:1-Sieg gegen den späteren Tabellenzweiten SC Goldau einen gelungenen Schlusspunkt unter die Saison.",
    "Mit einem lachenden und einem weinenden Auge blicken wir auf die vergangene Spielzeit zurück. Schweren Herzens verabschieden wir unser Trainerteam, Dani Schibli und Joelle Schibli. Für ihren unermüdlichen Einsatz, ihre wertvolle Unterstützung und ihr grosses Engagement danken wir ihnen von Herzen.",
    "Ebenso verabschieden wir uns von unseren Mitspielerinnen Nele Abel, Mia Sari, Geraldine Fedier, Xenia Krucker, Lejla Musliu, Tabea Gisler, Miren Brand und Jasmin Püntener. Vielen Dank für euren Einsatz, euren Teamgeist und die vielen gemeinsamen Erinnerungen.",
    "Mit dem Ende dieser Saison schliesst sich für das Team Uri Frauen II ein besonderes Kapitel. Gemeinsam werden wir uns dem Damen 1 anschliessen und freuen uns darauf, diese neue Herausforderung mit viel Motivation und Zusammenhalt anzugehen.",
]

FASNACHT = [
    "Auch in dieser Fasnacht konnten wir wieder die Katzenmusik Schattdorf sowie begeisterte Fasnächtler bei uns im Clubhaus begrüssen. Folgende Anlässe durften wir organisieren:",
    ("H3", "Freitag, 30. Januar 2026 — Vereinsfasnacht"),
    "Am Freitagabend konnten wir eine stattliche Anzahl Fasnächtler im Clubhaus begrüssen. Bei gemütlichem Zusammensein und der guten Unterhaltung von DJ Rämi konnten wir im Clubhaus einen gemütlichen Abend verbringen. Allen Helfern am Buffet und Grill ein herzliches Dankeschön für den grossartigen Abend.",
    ("H3", "Samstag, 31. Januar 2026 — SBU-Fasnacht"),
    "Am Samstag durften wir eine fröhliche Anzahl Fasnächtler auf dem Sportplatz Grüner Wald in Schattdorf begrüssen.",
    ("H3", "Samstag, 7. Februar 2026 — Jung & Alt Fasnacht"),
    "Auch am Samstag, 7. Februar 2026, durften wir wieder zusammen mit der KaMu Schattdorf einen grossartigen Anlass «Jung & Alt» organisieren. An dieser Stelle dem ganzen Vorstand der KaMu Schattdorf ein herzliches Dankeschön für ihre Unterstützung. Nebst dem Clubhaus und dem Grill sowie dem Unterhaltungsprogramm mit dem Künstler wurde den Anwesenden ein gemütlicher und unterhaltsamer Tag geboten. An dieser Stelle allen Helfern und Mitwirkenden ein herzliches Dankeschön.",
    ("H3", "Dankeschön für die «närrische» Zeit"),
    "Wir vom Clubhaus möchten uns herzlich bei allen Helfern und guten Seelen an den Anlässen im Clubhaus bedanken. Ohne die Unterstützung von vielen helfenden Händen wären solche Anlässe kaum zu bewältigen.",
]

JASS = [
    "Am Freitag, 16. Januar 2026, trafen sich 26 Paare (52 Teilnehmer) zum jährlichen Jassturnier im Clubhaus. Das in den Vorjahren bestens organisierte Turnier unter der Leitung von Mario Trovatelli ist ein bekannter Anlass im Jahresprogramm des FCS. Die diesjährige Organisation unter der neuen Leitung von Orlando Gisler, Simon Geisser, Reto Infanger und Mario Trovatelli kann sicher als Erfolg abgeschlossen werden. Nach dem Start konnten nach der 2. Runde alle Teilnehmer am Grill verpflegt werden, sodass der 2. Teil gestärkt in Angriff genommen werden konnte.",
    "Als Preise konnten wir den ersten 10 Jasspaaren ein kleines Präsent überreichen. An dieser Stelle herzlichen Dank an die Sponsoren: Imholz Sport, Schelbert Bauunternehmung, AXA Winterthur (René Gnos) und Pilatus Flugzeugwerke AG.",
    "Um 22.30 Uhr waren die Auswertungen gemacht und man konnte zur Rangverkündigung übergehen:",
    "1. Rang: Ruedi Geisser und Urs Indergand\n2. Rang: Karl Schilter jun. und Karl Schilter sen.\n3. Rang: Iwan Herger und Heinz Hürlimann",
    "Der Abend wurde zu später Stunde erfolgreich und gemütlich beendet, und im Januar 2027 wird voraussichtlich der 20. Jassabend des FC Schattdorf stattfinden.",
    "Ein grosser Dank dem Clubhausteam Yvonne und Walti Gerber, den Grilleuren Werni Aschwanden, Kusi Gisler und Markus «Müx» Indergand sowie allen Teilnehmern ein herzliches Dankeschön für die Fairness – und bis zum nächsten Mal.",
]

# Schlussrangliste 19. Jassabend — vollständig gemäss «Rangliste 2026.pdf»:
# (Rang, Paar, P1, R1, P2, R2, ZwRg2, P3, R3, ZwRg3, P4, R4, Total)
JASS_RANGLISTE = [
    (1, "Ruedi Geisser / Urs Indergand", 1056, 3, 961, 13, 5, 1004, 10, 4, 1155, 2, 4176),
    (2, "Kari Schilter jun. / Kari Schilter sen.", 980, 11, 1081, 2, 3, 1078, 3, 1, 1007, 8, 4146),
    (3, "Iwan Herger / Heinz Hürlimann", 892, 20, 1067, 3, 9, 842, 21, 16, 1266, 1, 4067),
    (4, "Luca Herger / Marcel Fedier", 1023, 6, 945, 14, 8, 1016, 9, 6, 983, 11, 3967),
    (5, "Linus Arnold / Raphael Arnold", 1033, 5, 1039, 6, 2, 1003, 11, 2, 877, 18, 3952),
    (6, "Marlen Tresch / Sepp Tresch", 848, 23, 1050, 5, 14, 1055, 4, 8, 992, 9, 3945),
    (7, "Andre Zgraggen / Sandra Zgraggen", 990, 8, 1007, 9, 6, 880, 17, 11, 1009, 7, 3886),
    (7, "Stefan Imholz / Daniel Geisser", 1074, 2, 1056, 4, 1, 881, 16, 5, 875, 19, 3886),
    (9, "Agi Tresch / Stefan Tresch", 903, 17, 1023, 7, 11, 1042, 6, 7, 909, 14, 3877),
    (10, "Ralph Bomatter / Kurt Briker", 1095, 1, 828, 23, 12, 896, 14, 14, 1012, 6, 3831),
    (11, "Bärti Walker / Paul Zurfluh", 925, 15, 997, 10, 13, 988, 13, 9, 892, 17, 3802),
    (12, "Chrigu Zgraggen / Toni Arnold", 983, 9, 995, 11, 7, 1083, 2, 3, 729, 25, 3790),
    (13, "Felix Zurfluh / Stefan Baumann", 789, 26, 902, 16, 25, 1191, 1, 10, 901, 15, 3783),
    (14, "Lisi Geisser / Heidi Tresch", 894, 19, 817, 24, 23, 1024, 7, 20, 1029, 5, 3764),
    (15, "Ruedi Planzer / Sepp Arnold", 901, 18, 861, 21, 20, 838, 22, 24, 1145, 3, 3745),
    (16, "Stefan Aschwanden / Bruno Aschwanden", 981, 10, 889, 17, 16, 999, 12, 12, 875, 19, 3744),
    (17, "Sandra Scheiber / Andrea Brücker", 1036, 4, 923, 15, 9, 801, 25, 18, 950, 12, 3710),
    (18, "Pia Gnos / Theres Welti", 851, 22, 874, 20, 22, 1046, 5, 17, 934, 13, 3705),
    (19, "Josef Bissig / Elisabeth Bissig", 927, 14, 1092, 1, 4, 806, 24, 13, 872, 21, 3697),
    (20, "Walti Gerber / Mario Trovatelli", 904, 16, 792, 26, 24, 860, 20, 25, 1132, 4, 3688),
    (20, "Beatrice Baumann / Isabelle Brücker", 861, 21, 1010, 8, 15, 829, 23, 22, 988, 10, 3688),
    (22, "Peter Gisler / Henry Euler", 992, 7, 877, 19, 17, 885, 15, 19, 855, 22, 3609),
    (23, "Julia Arnold / Colette Müller", 957, 13, 887, 18, 18, 868, 18, 21, 896, 16, 3608),
    (24, "Adi Geisser / Louise Zgraggen", 810, 25, 982, 12, 19, 1022, 8, 15, 618, 26, 3432),
    (25, "Reto Infanger / Patrik Stampfli", 959, 12, 803, 25, 20, 862, 19, 23, 739, 24, 3363),
    (26, "Manuel Gnos / Meinrad Epp", 828, 24, 834, 22, 26, 693, 26, 26, 752, 23, 3107),
]

FASNACHT_KEIN_BERICHT = [
    "Samstag, 14. Februar 2026 — Fasnachtssamstag",
    "Samstag, 21. Februar 2026 — Vorstandsausflug 2026",
    "Donnerstag, 4. Juni 2026 — Kick in One",
    "Samstag, 13. Juni 2026 — Saisonabschluss",
]

TERMINE = [
    ("21.08.2026", "93. Generalversammlung, 19.00 Uhr",
     "Grüner Wald / Clubhaus FC Schattdorf"),
    ("Dez. 2026 [PRÜFEN]", "Klausfeier", "Uristier-Saal, Altdorf"),
    ("Jan. 2027", "20. Jassabend (voraussichtlich)", "Clubhaus"),
    ("[PLATZHALTER]", "Weitere Anlässe — durch Chef Events zu ergänzen", ""),
]

JUNIOREN_TITEL = "Von unnützen Statistiken und Hard Rock"
JUNIOREN = [
    "Die Schattdorfer A- bis D-Junioren haben in der vergangenen Frühlingsrunde bei 36 Siegen im Durchschnitt 4,24 mehr Tore als ihre Gegner geschossen. Klar gibt es wesentlich aussagekräftigere Statistiken als diese! Natürlich lebt der Fussball von Zahlen, Tabellen und Resultaten. Doch geht es vor allem im Juniorenbereich um wesentlich mehr: nämlich um die fussballerische und persönliche Entwicklung der jungen Spielerinnen und Spieler. Siege sollen gemeinsam gefeiert werden, aber ebenso wichtig ist das Lernen, mit Niederlagen umzugehen. Jugendförderung entsteht nicht durch Tore, sondern durch Charakterbildung, Teamgeist und Freude am Spiel.",
    "Nicht optimal starteten die A-Junioren in die Frühlingsrunde. Nach Veränderungen im Kader und verletzungsbedingten Ausfällen musste sich die Mannschaft zunächst neu finden. Von Spiel zu Spiel steigerte sich das Team jedoch deutlich. Dank gewonnener Partien gegen Spitzenreiter Cham und einem abschliessenden 6:0 gegen Perlen-Buchrain beendeten die Schattdorfer die Saison in der höchsten Juniorenliga, der Youth League, auf dem hervorragenden 3. Rang.",
    "Nach diversen Mutationen im Winter war das Kader der Ba-Junioren personell eher knapp besetzt. Dies wirkte sich auf die Trainingspräsenz aus, und einige Trainings mussten aufgrund mangelnder Teilnehmerzahl abgesagt werden. Gegen vermeintlich schwächere Gegner tat sich das Team phasenweise schwer und liess einige Punkte liegen. Dennoch resultierte am Ende ein guter 4. Rang. Der FC Schattdorf bedankt sich an dieser Stelle beim scheidenden Trainer Lon Simonaj für dessen Einsatz.",
    "Anders präsentierte sich die Situation bei den Bb-Junioren. Hier war die Trainingspräsenz sehr hoch. In der Frühlingsrunde spielte die Mannschaft neu in der Gruppe mit Nidwaldner und Luzerner Teams — eine klar schwerere Gruppe, wie sich herausstellte. Die Bb-Junioren spielten nicht nur erfolgreich, sondern auch ausserordentlich fair. Sie sicherten sich den 3. Tabellenrang, und mit nur 3 Strafpunkten gehörten die Schattdorfer zu den fairsten Teams in dieser Gruppe.",
    "Ein treuer, wenn auch unerwünschter Begleiter der Ca-Mannschaft blieb die Verletzungshexe. Wie schon zuvor in der Herbstrunde fielen wichtige Teamstützen aus und schwächten die Mannschaft. In einer sehr starken Gruppe konnte sich das Team leider nicht behaupten, und der Abstieg in die 2. Stärkeklasse konnte nicht verhindert werden. Die meisten Spieler werden in der kommenden Saison in den beiden B-Teams zum Einsatz kommen.",
    "Vier Minuten vor Meisterschaftsende gelang den Cb-Junioren ein grosser Coup: Im Derby gegen Altdorf erzielten sie den Treffer zum 1:2-Auswärtssieg. Dadurch verbesserten sie sich noch vom 5. auf den ausgezeichneten 2. Tabellenrang.",
    "Das Da-Team entwickelte sich im Frühling zu einer «richtigen Einheit», einer eingespielten Mannschaft mit starkem Zusammenhalt. Dem Trainerteam ist es gelungen, den Kids die «Essenzen» für eine erfolgreiche Juniorenarbeit zu vermitteln. Dies zeigte sich auch in den Resultaten: 5 Siegen standen 3 Niederlagen und 1 Unentschieden gegenüber.",
    "Die Db-Junioren schienen in dieser Saison häufig nach dem Motto «alles oder nichts» gespielt zu haben. In der Frühlingsrunde gab es kein Unentschieden zu notieren, dafür 5 Siege und deren 3 Niederlagen.",
    "AC/DC: Die Dc-Junioren haben gerockt! 10:0 gegen die Kickers, 7:1 gegen Brunnen. Sins wurde mit 5:0 abgefertigt. Und wie bei Rockstars gab es auch mal einen Kater: eine 6:0-Niederlage in Ibach. «Das Team entwickelte sich sowohl sportlich als auch menschlich hervorragend», lobt das engagierte Trainerteam. Das Team blieb in den vergangenen 3 Jahren praktisch unverändert und hat den Aufstieg von der 3. in die 2. und dann in die 1. Stärkeklasse geschafft! Eine aussergewöhnliche Leistung, welche mit Fleiss, Teamgeist und Einsatzbereitschaft erreicht wurde. Das Torverhältnis von 53:17 spricht für sich. Hells Bells!",
    "Nach einer starken Herbstrunde wagten die Dd-Juniorinnen und -Junioren das Abenteuer in der höheren 2. Stärkeklasse. Bei diesem Entscheid war allen klar, dass der FCS-Statistiker wohl weniger Siege notieren darf. Tatsächlich gingen die meisten Partien verloren. Der gute Teamspirit war immer da, und darauf darf die ganze Mannschaft stolz sein! Aus Niederlagen lernt man bekanntlich mehr, und dies hat die Kids sowohl in fussballerischer als auch menschlicher Sicht klar weitergebracht. Die Mädels wechseln nun zum Frauenfussball FF14/FF17, und auch bei den Jungs gibt es in der Sommerpause viele Veränderungen.",
    "Beim De-Team handelt es sich wohl um die «prominenteste» Juniorenmannschaft des FC Schattdorf. Jedenfalls erschienen die Jungs regelmässig bzw. am meisten aller Juniorenmannschaften im Urner Wochenblatt, oft sogar mit Foto. Kein Wunder, denn über erfolgreiche Mannschaften wird gerne berichtet. Als eines der jüngsten Teams in der höchsten Elite-Kategorie liessen sie ihren Gegnern nicht viele Punkte übrig. Ein spielerisch und kämpferisch starkes Kollektiv zeichnet das FCS-De-Team aus.",
    "An einem Spiel der Df-Junioren wird es garantiert nie langweilig. Es handelt sich um Spiele mit Torgarantie: In 6 ausgewerteten Spielen sind 109 Tore gefallen, also 18,17 Goals pro Spiel! Leider schossen hier die Gegner weit mehr als die doppelte Anzahl Tore.",
    "Die Weiterentwicklung jedes einzelnen Spielers und jeder einzelnen Spielerin stand bei den älteren E-Junioren im Fokus. Beim Ea-Team wurde dieses Ziel klar erreicht und sogar mit zahlreichen Videobeweisen dokumentiert. Die meisten Spiele in der 1. Stärkeklasse konnten siegreich gestaltet werden. Das Turnier in Schwyz bleibt den Eb- und Ec-Teams sicher in Erinnerung: Sie dominierten ihre Gegner und gewannen alle ihre Partien. Bei den Ed/Ee-Junioren spielten vorwiegend die jüngeren Jahrgänge in der 2. Stärkeklasse. Vor allem im Format 3 gegen 3 zeigten die Schattdorfer spielerisch starke Leistungen und hatten ihre Gegner oft gut im Griff. Im 6-gegen-6-Spiel besteht noch Entwicklungspotenzial. Die Basis für die kommende Saison ist gelegt, und der FCS kann erneut mit zwei ausgeglichenen Teams antreten, künftig sogar in der 1. Stärkeklasse. Ohne Auswechselspieler trat die Ef-Equipe beim Turnier in Buochs an, und trotzdem gab es viel zu jubeln. Die jungen Fussballer zeigten hervorragenden Einsatz und gestalteten die meisten Spiele erfolgreich.",
    "Nicht ausschliesslich um Fussball geht es bei der «FF-Turngruppe». Zwischendurch zeigen die Girls auch gerne einen Überschlag oder andere turnerische Kunststücke. Mit dem Zuhören klappt es bei den Traineranweisungen nicht immer perfekt, denn die Gespräche untereinander sind oft in vollem Gange. Weibliches Multitasking eben! Viel wichtiger als die Resultate sind den Mädels die gemeinsamen Erlebnisse. Ein perfekter Turniertag gelang den Girls in Oberägeri, wo tatsächlich alle Spiele gewonnen wurden.",
    "Bei den F-Junioren geht es mit viel Energie und Freude zu und her. Die Kinder lernen von ihren Trainern mit Begeisterung neue Techniken – oder bringen sich diese gleich selber bei. Der Spass am Spiel steht immer im Vordergrund. Dabei spielt das Ausbildungsformat «play more football» mit dem Spiel auf kleine und grössere Tore eine wichtige Rolle in der Grundausbildung. Bei vielen der jüngsten Kicker ist bereits heute zu erkennen, dass sie auch in den höheren Kategorien für Furore sorgen können.",
    "Und zum Schluss noch etwas sehr Wichtiges: Wenn die Schattdorfer Junioren in der Frühlingsrunde verloren (insgesamt 43-mal), dann hat der Gegner im Durchschnitt 4,32 Tore mehr geschossen!",
]

FCS3_STAT = [
    ("Reto Bissig", 17, 17, 0, 0, 2),
    ("Miguel Letra Moisés", 16, 15, 1, 0, 2),
    ("Juri Spillmann (C)", 15, 14, 1, 2, 1),
    ("Robel Michael", 15, 10, 5, 0, 0),
    ("Silas Arnold", 14, 13, 1, 2, 8),
    ("Matteo Zberg", 13, 12, 1, 0, 0),
    ("Ranadan Fragnito", 13, 8, 5, 0, 0),
    ("Yannic Jäger", 12, 8, 4, 0, 0),
    ("Kay Schillig", 11, 11, 0, 1, 0),
    ("Livio Arnold", 11, 5, 6, 1, 1),
    ("Aleksandar Stojanovic", 11, 8, 3, 0, 1),
    ("Claudio Inderkum", 11, 11, 0, 1, 1),
    ("Mariano Prandi", 10, 8, 2, 0, 0),
    ("Patrick Gamma", 9, 2, 7, 0, 4),
    ("David Felix Rodrigues", 9, 6, 3, 0, 0),
    ("Claudio Pfyl", 7, 6, 1, 0, 0),
    ("Renato Rickli", 6, 6, 0, 0, 0),
    ("Renato José Da Fonseca Sobreira", 6, 6, 0, 1, 1),
    ("Josua Müller", 5, 5, 0, 1, 1),
    ("Luka Cota", 5, 5, 0, 2, 1),
    ("Ali Kocaslan", 4, 0, 4, 0, 0),
    ("Jannik Arnold", 3, 3, 0, 0, 0),
    ("Kamil Borkowski", 3, 1, 2, 1, 0),
    ("Pedro Neto Mulle", 2, 0, 2, 0, 0),
    ("Luca Herger", 2, 1, 1, 0, 1),
    ("Gian Luca Furger", 2, 2, 0, 0, 0),
    ("Joel Brand", 1, 1, 0, 1, 0),
    ("Mohammad Jawad Alizadeh", 1, 0, 1, 0, 0),
]

# «Wichtige Adressen» — 1:1 übernommen aus der Ausgabe Winter 2025/26
ADRESSEN = [
    ("Präsident", "Bomatter Ralph, Kirchgasse 1a, Schattdorf", "079 390 42 01"),
    ("Vizepräsident", "Herger Iwan, Klausenstrasse 11a, Altdorf", "078 698 91 95"),
    ("Marketing", "Schorno Patrick, Zwyergasse 27, Altdorf", "079 273 78 11"),
    ("Administration", "Deplazes Monja, Achern 54, Schattdorf", "078 658 44 34"),
    ("Finanzen", "Gisler Claudia, Stiege 32, Bürglen", "078 859 55 83"),
    ("Infrastrukturen", "Planzer Reto, Gandrütti 47, Schattdorf", "079 487 31 26"),
    ("Wettspielbetrieb", "Indergand Markus, Gotthardmatte 26, Schattdorf", "079 219 66 27"),
    ("Sportchef", "Gnos René, Grundmatte 9, Schattdorf", "079 420 61 20"),
    ("Junioren-Obmann", "Herger Iwan, Klausenstrasse 11a, 6460 Altdorf", "078 698 91 95"),
    ("Veranstaltungen", "Gisler Orlando, Wegmatt 16, Altdorf", "079 454 25 07"),
    ("Protokoll", "Küttel Jasmin, Langgasse 9, Schattdorf", "041 870 86 30"),
    ("Sponsoring", "Schorno Patrik, Zwyergasse 27, Schattdorf", "079 273 78 11"),
    ("Kommunikation", "Scheiber Dominique, Hintere Schilligmatte 4, Bürglen", "079 296 27 26"),
    ("Trainer 1. Ms.", "Piccirillo Emanuele, Parkpromenade 4, Emmen", "079 415 27 99"),
    ("Coach 1. Ms.", "Infanger Reto, Adlergartenstr. 11, Schattdorf", "079 531 83 22"),
    ("Trainer 2. Ms.", "Lussmann Mathias, Rathausplatz 3, Altdorf", "079 265 01 68"),
    ("Coach 2. Ms.", "Zurfluh Roger, Eyrütti 6, Schattdorf", "079 372 64 26"),
    ("Trainer 3. Ms.", "Jäger Yannic, Militärstrasse 12, Schattdorf", "079 173 36 05"),
    ("Trainer Frauen I", "vakant", "–"),
    ("Trainer Frauen II", "Schibli Daniel, Bahnhofstrasse 42, Altdorf", "079 514 57 57"),
    ("Coach Frauen II", "Schibli Joelle, Bahnhofstrasse 42, Altdorf", "078 944 76 66"),
    ("Verantwortliche Frauen", "Scheiber Dominique, Hintere Schilligmatte 4, Bürglen", "079 296 27 26"),
    ("Trainer Senioren", "Britschgi Marino, Gotthardmatte 4, Schattdorf", "079 637 54 98"),
    ("Verantwortlicher Sen.", "Britschgi Marino, Gotthardmatte 4, Schattdorf", "079 637 54 98"),
    ("«Schiri»-Verantw.", "Gisler Stephan, Attinghauserstrasse 43, Altdorf", "079 234 91 90"),
    ("J+S-Coach", "Arnold Karl, Furrersgrund 3, Altdorf", "079 159 82 38"),
    ("SFV-Anmeldestelle", "Hauger Roger, Bristenstrasse 9, Altdorf", "078 682 14 68"),
    ("OK-Chef Grümpelt.", "Arnold Paul, Bärengässli 5, Altdorf", "079 706 55 64"),
    ("Präsident Top Club", "Mahrow Kai, Bötzlingerstrasse 26, Schattdorf", "079 605 08 19"),
]


# ------------------------------------------------------------------ Statistik


def lade_ligatabellen():
    import openpyxl
    wb = openpyxl.load_workbook(
        os.path.join(SRC, "AktivTeams", "Statistik.xlsx"), data_only=True)
    tabellen = {}
    for name in ("FCS 1", "FCS2", "FCS 3", "Senioren ", "Frauen1", "Frauen 2"):
        ws = wb[name]
        zeilen = []
        titel = None
        for row in ws.iter_rows(values_only=True):
            vals = [v for v in row if v is not None and str(v).strip()]
            if not vals:
                continue
            if titel is None:
                titel = str(vals[0]).strip()
                continue
            # Rang, Verein, Sp, S, U, N, SP, Tore, :, Gegentore, Diff, Pkt
            r = [str(v).strip() if v is not None else "" for v in row]
            zeilen.append([r[0], r[1], r[2], r[3], r[4], r[5],
                           r[6].lstrip("-") if r[6] else "0",
                           f"{r[7]} : {r[9]}", r[11]])
        tabellen[name.strip()] = (titel, zeilen)
    # Einsatzstatistik FCS 1
    ws = wb["Statistik"]
    einsatz = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue
        einsatz.append([f"{row[0]} {row[1]}", row[2], row[3], row[4], row[5]])
    return tabellen, einsatz


def lade_fotoklassifikation():
    """Zeilen: DATEINAME | KATEGORIE | FORMAT | SCORE | BESCHREIBUNG"""
    if not os.path.exists(FOTO_KLASSIFIKATION):
        return []
    fotos = []
    for line in open(FOTO_KLASSIFIKATION, encoding="utf-8"):
        parts = [p.strip() for p in line.split("|")]
        if len(parts) < 4:
            continue
        pfad = os.path.join(SRC, "Fotos", parts[0])
        if not os.path.exists(pfad):
            continue
        try:
            score = int(parts[3])
        except ValueError:
            continue
        fotos.append({"pfad": pfad, "kategorie": parts[1],
                      "format": parts[2], "score": score})
    return fotos


def waehle_fotos(fotos, kategorien, minimum=4, limit=12):
    treffer = [f for f in fotos
               if f["kategorie"] in kategorien and f["score"] >= minimum]
    treffer.sort(key=lambda f: -f["score"])
    return [f["pfad"] for f in treffer[:limit]]


# ---------------------------------------------------------------------- Build


def main():
    doc = Document()

    # Grundstil
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    rpr = normal.element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    rpr.append(rfonts)

    # ---- Sektion 1: Titelseite (schmale Ränder, kein Kopf/Fuss)
    cover = doc.sections[0]
    cover.page_width, cover.page_height = Cm(21), Cm(29.7)
    for edge in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(cover, edge, Cm(0.7))
    titelbild = os.path.join(SRC, "Titelbild",
                             "Gemini_Generated_Image_io7s66io7s66io7s.png")
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(par)
    if os.path.exists(titelbild):
        cached, ratio = prep_image(titelbild, max_px=2400)
        par.add_run().add_picture(cached, height=Cm(26.2))
    band = doc.add_paragraph()
    _shade_paragraph(band, HEX_INK)
    band.alignment = WD_ALIGN_PARAGRAPH.CENTER
    band.paragraph_format.space_before = Pt(8)
    band.paragraph_format.space_after = Pt(8)
    run = band.add_run("SOMMER 2026  ·  FC SCHATTDORF  ·  "
                       "38. JAHRGANG, 74. NUMMER [PRÜFEN]")
    _set_font(run, size=11, bold=True, color=WHITE, spacing=30)

    # ---- Sektion 2: Inhalt (normale Ränder, Masthead)
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    build_masthead(section)

    # Inserate-Seite (wie in allen Ausgaben direkt nach dem Titelbild)
    add_centered_image(doc, os.path.join(INSERATE, "synaxis.png"),
                       width_cm=15.0, before=18, after=24)
    add_centered_image(doc, os.path.join(INSERATE, "enz.png"),
                       width_cm=12.5, before=24, after=6)
    add_inserat_hinweis(doc)

    # Inhaltsverzeichnis
    doc.add_page_break()
    add_h2(doc, "Inhalt", space_before=4)
    par = doc.add_paragraph()
    _no_space(par, after=4)
    run = par.add_run("Seitenzahlen aktualisieren: Verzeichnis anklicken und "
                      "F9 drücken (bzw. Rechtsklick → «Felder aktualisieren»).")
    _set_font(run, size=8, color=MUTED, italic=True)
    par = doc.add_paragraph()
    _field(par, 'TOC \\o "1-2" \\h \\z \\u')

    # Impressum, Vereinsdaten, Adressen — eigene Seite, 1:1 wie Winter 25/26
    doc.add_page_break()
    add_ressort_band(doc, "Impressum, Vereinsdaten, Adressen")
    impressum = [
        ("Herausgeber", "FC Schattdorf, Schattdorf, www.fcschattdorf.ch"),
        ("Redaktion", "Scheiber Dominique, Hintere Schilligmatte 4, Bürglen — 079 296 27 26"),
        ("Redaktionsschluss", "[PLATZHALTER]"),
        ("Inserat-Annahme", "FC Schattdorf, FCS-Zyttig, 6467 Schattdorf"),
        ("Auflage", "Online / 70 gedruckt"),
        ("Ausgabe", "Sommer 2026 / 38. Jahrgang, 74. Nummer [PRÜFEN]"),
    ]
    for label, wert in impressum:
        par = doc.add_paragraph()
        _no_space(par, after=1)
        run = par.add_run((label + ":").ljust(22))
        _set_font(run, size=9, bold=True, color=INK)
        run = par.add_run(wert)
        _set_font(run, size=9, color=INK)

    add_h3(doc, "Wichtige Adressen")
    widths = [Cm(4.6), Cm(9.4), Cm(3.4)]
    table = doc.add_table(rows=len(ADRESSEN), cols=3)
    _table_fixed(table, widths)
    for r, (funktion, person, tel) in enumerate(ADRESSEN):
        for i, val in enumerate((funktion, person, tel)):
            cell = table.rows[r].cells[i]
            _cell_margins(cell, top=14, bottom=14)
            _cell_borders(cell, ("bottom",))
            par = cell.paragraphs[0]
            _no_space(par)
            run = par.add_run(val)
            _set_font(run, size=8, bold=(i == 0), color=INK)

    add_h3(doc, "Vereinsdaten")
    for label, wert in [
        ("SFV Klub-Nr.", "02137"),
        ("UKB IBAN", "CH75 0078 001 9474 1"),
        ("1. Gründung", "1916"),
        ("2. Gründung", "1933"),
        ("Sportplatz", "Grüner Wald (100 × 61.5 m)"),
    ]:
        par = doc.add_paragraph()
        _no_space(par, after=1)
        run = par.add_run((label + ":").ljust(22))
        _set_font(run, size=9, bold=True, color=INK)
        run = par.add_run(wert)
        _set_font(run, size=9, color=INK)
    par = doc.add_paragraph()
    _no_space(par, before=4)
    run = par.add_run("[PRÜFEN] Adressen 1:1 aus der Ausgabe Winter 2025/26 "
                      "übernommen — Mutationen auf die Saison 2026/27 durch "
                      "die Administration nachführen.")
    _set_font(run, size=7.5, color=MUTED, italic=True)

    # ---- PRÄSIDENT
    doc.add_page_break()
    add_ressort_band(doc, "Präsident", "Ralph Bomatter")
    add_kicker(doc, "Vorwort")
    par = doc.add_paragraph()
    _no_space(par, after=8)
    run = par.add_run("VORWORT DES PRÄSIDENTEN")
    _set_font(run, size=20, bold=True, color=INK, spacing=15)
    par.style = doc.styles["Heading 2"]
    _restyle_heading(par, size=20)
    for absatz in VORWORT:
        add_body(doc, absatz, size=9.5, after=4, line=1.1)
    add_signature(doc, "Ralph Bomatter, Präsident")

    # GAMMA — kommt in jeder Ausgabe direkt nach dem Vorwort
    add_centered_image(doc, os.path.join(INSERATE, "gamma.png"),
                       width_cm=15.0, before=60, break_before=True)
    add_inserat_hinweis(doc)

    doc.add_page_break()
    add_kicker(doc, "Generalversammlung")
    h = add_h2(doc, "Einladung zur 93. Generalversammlung", space_before=2)
    h.style = doc.styles["Heading 2"]
    _restyle_heading(h)
    add_body(doc, "Freitag, 21. August 2026, 19.00 Uhr,", bold=True, after=0)
    add_body(doc, "auf dem «Grünen Wald» und im Clubhaus des FC Schattdorf",
             bold=True, after=10)
    add_h3(doc, "Traktandenliste 93. GV vom 21. August 2026")
    for nr, titel, subs in TRAKTANDEN:
        par = doc.add_paragraph()
        _no_space(par, after=1)
        par.paragraph_format.space_before = Pt(3)
        run = par.add_run(nr.ljust(5))
        _set_font(run, size=10, bold=True, color=RED_DEEP)
        run = par.add_run(titel)
        _set_font(run, size=10, bold=True, color=INK)
        for sub in subs:
            par = doc.add_paragraph()
            _no_space(par, after=1)
            par.paragraph_format.left_indent = Cm(1.0)
            run = par.add_run(sub)
            _set_font(run, size=9.5, color=INK)

    # ---- SPONSORING (Platzhalter)
    doc.add_page_break()
    add_ressort_band(doc, "Sponsoring", "Sponsoringverantwortlicher — "
                                        "wird extern bearbeitet")
    add_placeholder_box(doc, "Sponsoren, die die FCS-Zyttig gedruckt erhalten sollen", [
        "Liste der Sponsoren einfügen, die ein gedrucktes Exemplar erhalten.",
    ])
    add_placeholder_box(doc, "Sponsoren, die erstmals in der FCS-Zyttig sponsoren", [
        "Neue Sponsoren mit Logo/Inserat aufführen und willkommen heissen.",
    ])
    add_placeholder_box(doc, "Kontrolle bisheriger Inserate", [
        "Bestehende Inserate prüfen: noch aktuell? Vertrag verlängert?",
        "Inserate-Seiten an den gewünschten Stellen im Dokument einfügen.",
    ])

    # ---- ADMINISTRATION (Platzhalter + Adressliste)
    doc.add_page_break()
    add_ressort_band(doc, "Administration", "Chefin Administration "
                                            "Monja Deplazes — wird extern bearbeitet")
    add_placeholder_box(doc, "Inhalte Administration", [
        "Mutationen / Mitgliederwesen, Hinweise zu Mitgliederbeiträgen,",
        "weitere administrative Mitteilungen — durch Administration zu ergänzen.",
        "(Die wichtigen Adressen stehen neu auf der Impressum-Seite.)",
    ])

    # ---- EVENTS
    doc.add_page_break()
    add_ressort_band(doc, "Events", "Chef Veranstaltungen Orlando Gisler")
    fotos = lade_fotoklassifikation()

    fasnacht_bilder = extrahiere_fasnacht_bilder()
    add_kicker(doc, "Rückblick · Fasnacht")
    h = add_h2(doc, "Fasnachtsanlässe 2026 im Clubhaus", space_before=2)
    h.style = doc.styles["Heading 2"]
    _restyle_heading(h)
    add_body(doc, FASNACHT[0])
    # Anlässe ohne Bericht (gemäss Hinweis des Verfassers)
    table = doc.add_table(rows=1, cols=1)
    _table_fixed(table, [CONTENT_W])
    cell = table.rows[0].cells[0]
    _shade_cell(cell, HEX_MIST)
    _cell_borders(cell, ("start",), hexcolor=HEX_RED, size=24)
    _cell_margins(cell, top=90, bottom=90, left=180, right=180)
    par = cell.paragraphs[0]
    _no_space(par, after=2)
    run = par.add_run("Von folgenden Anlässen folgt kein Bericht:")
    _set_font(run, size=9, bold=True, color=INK)
    for zeile in FASNACHT_KEIN_BERICHT:
        par = cell.add_paragraph()
        _no_space(par, after=1)
        run = par.add_run("•  " + zeile)
        _set_font(run, size=9, color=MUTED)
    doc.add_paragraph()
    # Abschnitte mit den Bildern aus dem eingereichten Bericht
    add_h3(doc, FASNACHT[1][1])
    add_body(doc, FASNACHT[2])
    add_photo_grid(doc, fasnacht_bilder["vereinsfasnacht"], cols=3,
                   max_h_cm=5.2)
    add_h3(doc, FASNACHT[3][1])
    add_body(doc, FASNACHT[4])
    add_photo_grid(doc, fasnacht_bilder["sbu"], cols=3, max_h_cm=5.2)
    add_h3(doc, FASNACHT[5][1])
    add_body(doc, FASNACHT[6])
    add_photo_grid(doc, fasnacht_bilder["jungalt"], cols=3, max_h_cm=5.2)
    add_h3(doc, FASNACHT[7][1])
    add_body(doc, FASNACHT[8])
    add_photo_grid(doc, fasnacht_bilder["dank"], cols=2, max_h_cm=6.5)
    add_signature(doc, "Orlando Gisler")

    doc.add_page_break()
    add_kicker(doc, "Rückblick · Jassen")
    h = add_h2(doc, "19. Jassturnier des FC Schattdorf", space_before=2)
    h.style = doc.styles["Heading 2"]
    _restyle_heading(h)
    _render_mixed(doc, JASS)
    add_stat_tabelle(
        doc, "Schlussrangliste — 19. Jassabend, 16. Januar 2026",
        ["Rg", "Jasspaar", "P1", "Rg", "P2", "Rg", "ZR", "P3", "Rg",
         "ZR", "P4", "Rg", "Total"],
        [list(z) for z in JASS_RANGLISTE],
        [Cm(0.9), Cm(4.9), Cm(1.0), Cm(0.8), Cm(1.0), Cm(0.8), Cm(0.95),
         Cm(1.0), Cm(0.8), Cm(0.95), Cm(1.0), Cm(0.8), Cm(1.3)],
        legende="P1–P4 Punkte pro Passe · Rg Rang der jeweiligen Passe · "
                "ZR Zwischenrang nach der 2. bzw. 3. Passe · "
                "Quelle: offizielle Rangliste")

    add_h2(doc, "Events in Bildern")
    event_fotos = waehle_fotos(fotos, {"EVENT_FEST"}, minimum=3, limit=8)
    if event_fotos:
        add_body(doc, "[PRÜFEN] Bildauswahl/Zuordnung durch Utsch oder Silvana "
                      "bestätigen lassen.", size=8.5, color=MUTED, italic=True)
        add_photo_grid(doc, event_fotos, cols=2)
    add_placeholder_box(doc, "Fotos von Events", [
        "Weitere Event-Bilder von Utsch oder Silvana einfügen (sofern vorhanden).",
    ])

    add_h2(doc, "Bevorstehende Veranstaltungen")
    add_agenda(doc, TERMINE)
    doc.add_paragraph()

    # ---- SPORTCHEF
    tabellen, einsatz = lade_ligatabellen()
    doc.add_page_break()
    add_ressort_band(doc, "Sportchef", "René Gnos")

    def team_artikel(kicker, titel, absaetze, signatur, tabelle_key,
                     highlight, sponsoren, bilder=(), extra=None,
                     seitenumbruch=True, logos=None):
        if seitenumbruch:
            doc.add_page_break()
        add_kicker(doc, kicker)
        h = add_h2(doc, titel, space_before=2)
        h.style = doc.styles["Heading 2"]
        _restyle_heading(h)
        _render_mixed(doc, absaetze)
        if signatur:
            add_signature(doc, signatur)
        bilder = [os.path.join(SRC, "AktivTeams", b) for b in bilder]
        add_photo_grid(doc, bilder, cols=2 if len(bilder) > 1 else 1)
        if tabelle_key:
            titel_t, zeilen = tabellen[tabelle_key]
            add_liga_tabelle(doc, "Schlusstabelle " + titel_t, zeilen, highlight)
        if extra:
            extra()
        if sponsoren:
            add_sponsor_box(doc, sponsoren, logos=logos)

    def fcs1_einsatz():
        add_stat_tabelle(
            doc, "Einsatzstatistik 1. Mannschaft",
            ["Spieler", "Einsätze", "Minuten", "Tore", "Gelb"],
            einsatz, [Cm(6.6), Cm(2.7), Cm(2.7), Cm(2.7), Cm(2.7)],
            legende="Gelb = Verwarnungen")

    def fcs3_stat():
        add_stat_tabelle(
            doc, "Teaminterne Statistik FCS 3",
            ["Spieler", "Einsätze", "Startelf", "Joker", "Gelb", "Tore"],
            FCS3_STAT, [Cm(7.4), Cm(2.0), Cm(2.0), Cm(2.0), Cm(2.0), Cm(2.0)],
            legende="Joker = Einwechslungen · ausgewertete Spiele: 18 · "
                    "Instagram: @fcs_3_officialpage")

    team_artikel(
        "2. Liga regional · 1. Mannschaft",
        "Herausfordernd, lehrreich und immer im Sinne des FC Schattdorf",
        FCS1, "Sportliche Grüsse — René Gnos, Sportchef FC Schattdorf",
        "FCS 1", "Schattdorf",
        [("Team-Sponsoren", "Wohncenter Muoser Schattdorf · Imholz Sport "
          "Bürglen · Gasthaus Brückli Schattdorf · Kebab Hüsli Schattdorf")],
        bilder=("FCS1.JPG", "FCS1-Abschied.JPG"),
        extra=fcs1_einsatz, seitenumbruch=False,
        logos=[(os.path.join(INSERATE, "gasthaus-brueckli.png"), 7.0)])

    # Mobiliar — fester Platz bei der 1. Mannschaft
    add_centered_image(doc, os.path.join(INSERATE, "mobiliar.png"),
                       width_cm=11.5, before=10)
    add_inserat_hinweis(doc)

    team_artikel(
        "4. Liga · 2. Mannschaft",
        "Rückblick auf die Saison der 2. Mannschaft",
        FCS2, "Die 2. Mannschaft",
        "FCS2", "Schattdorf",
        [("Dress-Sponsor", "Gasthaus Brückli, Schattdorf")],
        bilder=("FCS 2_Web.jpg", "FCS 2 DerbySieg.jpeg"))

    # Druckerei Kuster — fester Platz bei der 2. Mannschaft
    add_centered_image(doc, os.path.join(INSERATE, "kuster.png"),
                       width_cm=11.5, before=10)
    add_inserat_hinweis(doc)

    # Inserate-Seite bei der 2. Mannschaft (fixer Platz in jeder Ausgabe)
    doc.add_page_break()
    add_centered_image(doc, os.path.join(INSERATE, "brand-automobile.png"),
                       width_cm=14.5, before=12, after=18)
    add_centered_image(doc, os.path.join(INSERATE, "schibli.png"),
                       width_cm=11.5, before=18, after=6)
    add_inserat_hinweis(doc)

    team_artikel(
        "5. Liga · 3. Mannschaft",
        "Torhüter-Chaos in der 5. Liga",
        FCS3, None,
        "FCS 3", "Schattdorf",
        [("Dress-Sponsor", "BINARY one GmbH")],
        bilder=("FCS 3.jpg",),
        extra=fcs3_stat)

    # Inserate-Seite bei der 3. Mannschaft (fixer Platz in jeder Ausgabe)
    doc.add_page_break()
    add_centered_image(doc, os.path.join(INSERATE, "urner-kantonalbank.png"),
                       width_cm=13.0, before=12, after=16)
    add_centered_image(doc, os.path.join(INSERATE, "brand-maison-metall.png"),
                       width_cm=9.5, before=16, after=6)
    add_inserat_hinweis(doc)

    team_artikel(
        "Senioren 30+ Promotion · Team Uri",
        "Rückrundenbericht 2026 — jetzt mit echter Fussballhumor-DNA",
        SENIOREN, "Senioren Team Uri",
        "Senioren", "Team Uri",
        [("Team-Sponsoren", "Herger Küchen AG · Mövenpick · "
          "Maler Nideröst AG · Linden Apotheke")],
        logos=[(os.path.join(INSERATE, "herger-kuechen.png"), 7.0)])

    # Inserate-Seite bei den Senioren (fixer Platz in jeder Ausgabe)
    doc.add_page_break()
    add_centered_image(doc, os.path.join(INSERATE, "arnold-ag.png"),
                       width_cm=14.5, before=12, after=18)
    add_centered_image(doc, os.path.join(INSERATE, "eichhof.png"),
                       width_cm=12.5, before=18, after=6)
    add_inserat_hinweis(doc)

    team_artikel(
        "Frauen 2. Liga · Team Uri 1",
        "Saisonbericht Team Uri Frauen 1",
        FRAUEN1, "Team Uri Frauen 1",
        "Frauen1", "Team Uri",
        [("Dress-Sponsor", "Raiffeisen")],
        bilder=("Team Uri 1 Teamfoto.jpeg", "Cupfinale Frauen Team Uri 1.jpeg",
                "Team Uri 1 Einzug Cup Finale.jpeg", "Team Uri 1 Cup Finale.jpeg"))

    team_artikel(
        "Frauen 4. Liga · Team Uri 2",
        "Saisonbericht Team Uri Frauen 2",
        FRAUEN2, "Team Uri Frauen 2",
        "Frauen 2", "Team Uri",
        [("Dress-Sponsor", "TEKO")],
        bilder=("Team Uri 2 Teamfoto.jpeg", "Team Uri 2 Siegesfoto.jpeg"))

    # ---- JUNIOREN
    doc.add_page_break()
    add_ressort_band(doc, "Junioren", "Juniorenobmann Iwan Herger")
    add_kicker(doc, "Rückblick Frühlingsrunde 2026")
    h = add_h2(doc, JUNIOREN_TITEL, space_before=2)
    h.style = doc.styles["Heading 2"]
    _restyle_heading(h)
    for absatz in JUNIOREN:
        add_body(doc, absatz)
    add_signature(doc, "Juniorenpresse: Linus Epp")

    junior_fotos = waehle_fotos(
        fotos, {"JUNIOREN_SPIEL", "JUNIOREN_LAGER"}, minimum=4, limit=20)
    if junior_fotos:
        doc.add_page_break()
        add_h2(doc, "Die Juniorenabteilung in Bildern", space_before=2)
        quer = [p for p in junior_fotos if _ist_quer(p)]
        hoch = [p for p in junior_fotos if not _ist_quer(p)]
        add_photo_grid(doc, quer, cols=2)
        add_photo_grid(doc, hoch, cols=3)
    else:
        add_placeholder_box(doc, "Fotos Juniorenabteilung", [
            "Bilder der Juniorenabteilung einfügen (sofern vorhanden).",
        ])

    add_sponsor_box(doc, [
        ("Junioren-Nachwuchs-Patronat", "GAMMA AG"),
        ("Weitere Sponsoren Juniorenabteilung",
         "[PLATZHALTER — durch Sponsoringverantwortlichen zu ergänzen]"),
    ], logos=[(os.path.join(INSERATE, "gamma.png"), 5.5)])

    # ---- SONSTIGES: alle noch nicht verwendeten Bilder
    verwendet = set(junior_fotos) | set(event_fotos)
    verwendet |= {os.path.join(SRC, "AktivTeams", b) for b in (
        "FCS1.JPG", "FCS1-Abschied.JPG", "FCS 2_Web.jpg",
        "FCS 2 DerbySieg.jpeg", "FCS 3.jpg",
        "Team Uri 1 Teamfoto.jpeg", "Cupfinale Frauen Team Uri 1.jpeg",
        "Team Uri 1 Einzug Cup Finale.jpeg", "Team Uri 1 Cup Finale.jpeg",
        "Team Uri 2 Teamfoto.jpeg", "Team Uri 2 Siegesfoto.jpeg")}
    alle_bilder = []
    for ordner in ("Fotos", "AktivTeams"):
        basis = os.path.join(SRC, ordner)
        for name in sorted(os.listdir(basis)):
            if name.lower().endswith((".jpg", ".jpeg", ".png")):
                alle_bilder.append(os.path.join(basis, name))
    rest = [p for p in alle_bilder if p not in verwendet]
    doc.add_page_break()
    add_ressort_band(doc, "Sonstiges", "Weitere Bilder aus dem Vereinsjahr — "
                                       "nicht verwendetes Material")
    quer = [p for p in rest if _ist_quer(p)]
    hoch = [p for p in rest if not _ist_quer(p)]
    add_photo_grid(doc, quer, cols=3, max_h_cm=4.4)
    add_photo_grid(doc, hoch, cols=4, max_h_cm=5.6)

    # ---- SWISSLIGHT (fixer Platz vor der Fussballschule)
    doc.add_page_break()
    add_centered_image(doc, os.path.join(INSERATE, "swisslight.png"),
                       width_cm=13.5, before=48)
    add_inserat_hinweis(doc)

    # ---- FUSSBALLSCHULE (aktuelle Flyer Herbst 2026)
    doc.add_page_break()
    add_ressort_band(doc, "Fussballschule")
    add_centered_image(doc, os.path.join(INSERATE,
                                         "fussballschule-herbst26.png"),
                       height_cm=21.5, before=6)
    doc.add_page_break()
    add_centered_image(doc, os.path.join(INSERATE,
                                         "schnuppertraining-herbst26.png"),
                       height_cm=23.5, before=6)

    # ---- TRAINER GESUCHT
    doc.add_page_break()
    add_ressort_band(doc, "Trainer gesucht")
    add_centered_image(doc, os.path.join(INSERATE, "trainer-gesucht-flyer.png"),
                       height_cm=21.0, before=6)

    # ---- Schlussseiten: MUOSER, dann Rückseite mit URNER TOR
    doc.add_page_break()
    add_centered_image(doc, os.path.join(INSERATE, "muoser.png"),
                       height_cm=23.0, before=12)
    add_inserat_hinweis(doc)
    doc.add_page_break()
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    par.add_run().add_picture(os.path.join(INSERATE, "pp-post.png"),
                              width=Cm(5.0))
    add_centered_image(doc, os.path.join(INSERATE, "urnertor.png"),
                       width_cm=13.5, before=140)
    add_inserat_hinweis(doc)

    doc.save(OUT_DOCX)
    print("OK:", OUT_DOCX)


def _ist_quer(path):
    img = ImageOps.exif_transpose(Image.open(path))
    return img.width >= img.height


def _restyle_heading(par, size=14.5):
    """Heading-Style zuweisen (für TOC), aber FCS-Look erzwingen."""
    for run in par.runs:
        _set_font(run, size=size, bold=True, color=INK, spacing=10)
    ppr = par._p.get_or_add_pPr()
    # Word-Standard-Heading-Farben/Fonts überschreiben passiert über Runs.


def _render_mixed(doc, bloecke):
    """Liste aus Strings und ('H3', titel)-Tupeln rendern."""
    for block in bloecke:
        if isinstance(block, tuple) and block[0] == "H3":
            add_h3(doc, block[1])
        else:
            add_body(doc, block)


if __name__ == "__main__":
    sys.exit(main())
