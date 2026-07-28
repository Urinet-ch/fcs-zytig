# -*- coding: utf-8 -*-
"""FCS-Zytig Sommer 2026 — Schlussbearbeitung nach der Korrekturrunde Bomatter.

Basis ist «Zytig Sommer 26/FCS-Zyttig Sommer 2026 ENTWURF - rbo.docx»
(vom Präsidenten im Korrekturmodus überarbeitete Fassung unseres Entwurfs).

Das Skript
  1. nimmt alle Korrekturmodus-Änderungen an (siehe docx_revisions.py),
  2. ersetzt das pixelige FCS-Wappen auf der Titelseite durch das Vektor-
     Wappen (fcs-logo-wappen.svg: Website-SVG, in dem die Spielerfigur weiss
     hinterlegt ist — hier auf Schwarz gesetzt, wie im gedruckten Wappen),
  3. tauscht Teamfoto und Dress-Sponsor der 3. Mannschaft (neu Feritec AG),
  4. setzt in jede Sponsorenbox die Farblogos aller genannten Sponsoren
     (assets/img/sponsoren, Herkunft siehe QUELLEN.txt dort),
  5. bestückt «Events in Bildern» mit Fotos vom 33. Dorf- und
     66. Grümpelturnier (Quelle: fcschattdorf.ch),
  6. entfernt die redaktionellen [PRÜFEN]/[PLATZHALTER]-Hinweise,
  7. hält Tabellen zusammen (kein Umbruch mitten in einer Tabelle),
  8. füllt die Fotoseiten so auf, dass die Seitenzahl durch 4 teilbar ist.

Aufruf:  python3 finalize_zytig.py [--fueller N]
"""

import copy
import hashlib
import io
import json
import os
import re
import subprocess
import sys

import docx
import docx.opc.packuri
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Emu, Pt
from lxml import etree
from PIL import Image, ImageDraw, ImageOps

import build_zytig_docx as bz
from docx_revisions import accept_revisions

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"

# Datumsformate der Termin-Agenda («21.08.2026», «Dez. 2026»)
DATUM = re.compile(r"^\s*(\d{1,2}\.\d{1,2}\.\d{4}|[A-Za-zÄÖÜäöü]{3,4}\.\s*\d{4})")

REPO = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(REPO, "Zytig Sommer 26")
INSERATE = os.path.join(REPO, "assets", "img", "inserate")
QUELLE = os.path.join(SRC, "FCS-Zyttig Sommer 2026 ENTWURF - rbo.docx")
ZIEL = os.path.join(REPO, "FCS-Zytig Sommer 2026 ENTWURF.docx")
GRUEMPI = os.path.join(SRC, "Fotos", "Gruempi")

# Grümpi-Bilder in Reihenfolge der Verwendung (Auswahl aus dem Homepage-Album)
GRUEMPI_QUER = [
    "thumbnail-13.jpg",    # Luftaufnahme Turniergelände
    "gymi-106.jpg",        # Torschuss vor Publikum
    "thumbnail-12.jpg",    # Spielszene vor dem Tor
    "thumbnail-5.jpg",     # jubelndes Damenteam
    "thumbnail-2.jpg",     # zwei Juniorinnen umarmen sich
    "thumbnail-8.jpg",     # Planschbecken zur Abkühlung
    "photo-2026-06-25-21-40-55.jpg",  # Siegerteam mit Pokal
    "thumbnail.jpg",       # Turnierbetrieb mit Zuschauern
]
GRUEMPI_HOCH = [
    "photo-2026-06-25-10-58-40.jpg",   # Siegerinnen-Team
    "photo-2026-06-25-21-36-48.jpg",   # Rangverkündigung am Abend
    "photo-2026-06-25-21-23-19.jpg",   # Festbetrieb im Zelt
]
# Sponsoren je Mannschaft — Logos in assets/img/sponsoren (Quellen: QUELLEN.txt).
# None = kein aktuelles Logo auffindbar, der Name steht nur im Text.
SPONSOREN = {
    # 1. Mannschaft — Team-Sponsoren gemäss fcschattdorf.ch/aktive/1-mannschaft
    "Wohncenter Muoser": {
        "text": [("Team-Sponsoren",
                  "Wohncenter Muoser Schattdorf · Imholz Sport Bürglen · "
                  "Gasthaus Brückli Schattdorf · Kebab Hüsli Schattdorf · "
                  "Axanova · Schelbert AG")],
        "logos": [("Wohncenter Muoser", "muoser.png"),
                  ("Imholz Sport", "imholz-sport.png"),
                  ("Gasthaus Brückli", "gasthaus-brueckli.png"),
                  ("Kebab Hüsli", "kebab-huesli.png"),
                  ("Axanova", "axanova.png"),
                  ("Schelbert AG", "schelbert.png")],
    },
    # 2. Mannschaft — Dress-Sponsor aus der Zytig, Team-Sponsor von der Website
    "Dress-Sponsor: Gasthaus Brückli": {
        "text": [("Team-Sponsor", "PORR")],
        "logos": [("Gasthaus Brückli", "gasthaus-brueckli.png"),
                  ("PORR", "porr.png")],
    },
    "Dress-Sponsor: Feritec": {
        "logos": [("Feritec AG", "feritec.png")],
    },
    # Senioren — Mövenpick und Lindenapotheke bleiben ohne Logo (Absprache)
    "Herger Küchen AG": {
        "logos": [("Herger Küchen AG", "herger-kuechen.png"),
                  ("Maler Nideröst AG", "maler-nideroest.png")],
    },
    "Dress-Sponsor: Raiffeisen": {
        "logos": [("Raiffeisen", "raiffeisen.png")],
    },
    "Dress-Sponsor: TEKO": {
        "text": [("Dress-Sponsor", "TEKO Oberflächentechnik AG")],
        "logos": [("TEKO Oberflächentechnik AG", "teko.png")],
    },
    "Junioren-Nachwuchs-Patronat": {
        "logos": [("GAMMA AG", "gamma.png")],
    },
}

# Reserve für den Seitenausgleich (Ressort «Sonstiges»)
GRUEMPI_RESERVE = [
    "thumbnail-6.jpg", "thumbnail-7.jpg", "thumbnail-10.jpg",
    "thumbnail-11.jpg", "thumbnail-9.jpg", "thumbnail-4.jpg",
    "thumbnail-3.jpg", "gymi-119.jpg", "thumbnail-15.jpg",
    "photo-2026-06-25-11-00-26.jpg", "photo-2026-06-25-10-59-34.jpg",
    "photo-2026-06-25-10-59-53.jpg",
]


# ------------------------------------------------------------------ Werkzeuge


def _md5(daten):
    return hashlib.md5(daten).hexdigest()


def bildteil_suchen(doc, quellpfad):
    """Bildteil im Dokument über den Byte-Hash der Quelldatei finden."""
    ziel = _md5(open(quellpfad, "rb").read())
    for rid, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue
        if _md5(rel.target_part.blob) == ziel:
            return rid, rel.target_part
    return None, None


def _als_bytes(bild, format_):
    puffer = io.BytesIO()
    if format_ == "JPEG":
        bild.convert("RGB").save(puffer, "JPEG", quality=88)
    else:
        bild.save(puffer, "PNG")
    return puffer.getvalue()


def bild_ersetzen(doc, rid, teil, neu, breite_cm=None, hoehe_cm=None,
                  hoehe_behalten=False):
    """Bildbytes austauschen und die Anzeigegrösse ans Seitenverhältnis anpassen.

    breite_cm/hoehe_cm: neue Anzeigegrösse; ohne Angabe bleibt die Breite
    bestehen und die Höhe wird aus dem neuen Seitenverhältnis berechnet.
    """
    format_ = "PNG" if teil.partname.ext.lower() == "png" else "JPEG"
    teil._blob = _als_bytes(neu, format_)
    ratio = neu.width / neu.height
    for blip in doc.element.body.iter(A + "blip"):
        if blip.get(R + "embed") != rid:
            continue
        rahmen = blip
        while rahmen is not None and rahmen.tag not in (WP + "inline",
                                                        WP + "anchor"):
            rahmen = rahmen.getparent()
        if rahmen is None:
            continue
        extent = rahmen.find(WP + "extent")
        cx = int(extent.get("cx"))
        cy = int(extent.get("cy"))
        if hoehe_behalten:
            cx = int(round(cy * ratio))
        else:
            if breite_cm:
                cx = int(Cm(breite_cm))
            cy = int(Cm(hoehe_cm)) if hoehe_cm else int(round(cx / ratio))
        extent.set("cx", str(cx))
        extent.set("cy", str(cy))
        for ext in rahmen.iter(A + "ext"):
            ext.set("cx", str(cx))
            ext.set("cy", str(cy))


def logo_rendern(svg, breite=1400):
    """SVG über die macOS-Vorschau scharf rendern, Hintergrund freistellen.

    qlmanage legt das Logo auf Weiss; das Wappen steht auf der Titelseite
    aber über dem Foto. Der weisse Rand wird deshalb von den Ecken her
    ausgegossen (die weisse Schrift im roten Balken bleibt erhalten).
    """
    ziel = os.path.join(bz.IMG_CACHE, "fcs-logo-%d.png" % breite)
    os.makedirs(bz.IMG_CACHE, exist_ok=True)
    if not os.path.exists(ziel):
        tmp = os.path.join(bz.IMG_CACHE, os.path.basename(svg) + ".png")
        subprocess.check_call(
            ["qlmanage", "-t", "-s", str(breite * 2), "-o", bz.IMG_CACHE, svg],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        bild = Image.open(tmp).convert("RGBA")
        hintergrund = Image.new("RGB", bild.size, "white")
        hintergrund.paste(bild, mask=bild.split()[3])
        marke = (255, 0, 255)
        for ecke in ((0, 0), (bild.width - 1, 0), (0, bild.height - 1),
                     (bild.width - 1, bild.height - 1)):
            ImageDraw.floodfill(hintergrund, ecke, marke, thresh=40)
        daten = hintergrund.convert("RGBA")
        pixel = daten.load()
        for y in range(daten.height):
            for x in range(daten.width):
                if pixel[x, y][:3] == marke:
                    pixel[x, y] = (255, 255, 255, 0)
        daten = daten.crop(daten.split()[3].getbbox())
        daten.thumbnail((breite, breite), Image.LANCZOS)
        daten.save(ziel)
        os.remove(tmp)
    return Image.open(ziel).convert("RGBA")



# ------------------------------------------------------- Absätze und Tabellen


def absatztext(el):
    return "".join(t.text or "" for t in el.iter(W + "t"))


def body_index(body, pruefer):
    for i, el in enumerate(body):
        if pruefer(el):
            return i
    return -1


def entfernen(el):
    el.getparent().remove(el)


def run_entfernen(body, teiltext):
    """Einzelnen Run mit dem Suchtext aus allen Absätzen löschen. → Anzahl."""
    n = 0
    for run in list(body.iter(W + "r")):
        text = "".join(t.text or "" for t in run.iter(W + "t"))
        if teiltext in text:
            entfernen(run)
            n += 1
    return n


def absaetze_entfernen(body, praefix):
    """Alle Absätze löschen, deren Text mit praefix beginnt. → Anzahl."""
    n = 0
    for par in list(body.iter(W + "p")):
        if absatztext(par).strip().startswith(praefix):
            entfernen(par)
            n += 1
    return n


def text_ersetzen(body, alt, neu):
    n = 0
    for t in body.iter(W + "t"):
        if t.text and alt in t.text:
            t.text = t.text.replace(alt, neu)
            n += 1
    return n


def neue_elemente(doc, aufbau):
    """Bausteine am Dokumentende erzeugen und wieder herauslösen."""
    body = doc.element.body
    vorher = list(body)
    aufbau()
    neu = [el for el in body if el not in vorher]
    for el in neu:
        body.remove(el)
    return neu


def einsetzen(body, index, elemente):
    for versatz, el in enumerate(elemente):
        body.insert(index + versatz, el)
    return index + len(elemente)


# ------------------------------------------------------------------ Bausteine


def gruempi_pfade(namen):
    return [os.path.join(GRUEMPI, n) for n in namen]


def events_in_bildern(doc):
    """Leeren Fotoblock durch die Grümpi-Bilder ersetzen."""
    body = doc.element.body
    kopf = body_index(body, lambda el: el.tag == W + "p"
                      and absatztext(el).strip() == "Events in Bildern")
    if kopf < 0:
        raise RuntimeError("Abschnitt «Events in Bildern» nicht gefunden")

    # Alles zwischen Titel und «Bevorstehende Veranstaltungen» wegräumen
    ende = body_index(body, lambda el: el.tag == W + "p"
                      and absatztext(el).strip() == "Bevorstehende Veranstaltungen")
    for el in list(body)[kopf + 1:ende]:
        entfernen(el)

    def aufbau():
        bz.add_body(doc, "33. Dorf- und 66. Grümpelturnier vom 18. bis "
                         "20. Juni 2026 auf dem «Grünen Wald».",
                    size=9.5, after=8)
        bz.add_photo_grid(doc, gruempi_pfade(GRUEMPI_QUER), cols=2,
                          max_h_cm=6.4)
        bz.add_photo_grid(doc, gruempi_pfade(GRUEMPI_HOCH), cols=3,
                          max_h_cm=7.0)

    einsetzen(body, kopf + 1, neue_elemente(doc, aufbau))


def fueller_seiten(doc, anzahl, pro_seite=6):
    """Ganze Fotoseiten anhängen, damit die Seitenzahl durch 4 teilbar wird.

    Die Seiten kommen ans Ende des Ressorts «Sonstiges», also vor die
    Inserate- und Flyer-Seiten am Schluss des Hefts.
    """
    if anzahl <= 0:
        return
    body = doc.element.body
    start = body_index(body, lambda el: el.tag == W + "p"
                       and absatztext(el).strip() == "Sonstiges")
    if start < 0:
        raise RuntimeError("Ressort «Sonstiges» nicht gefunden")
    ziel = start + 1
    for i, el in enumerate(list(body)[start:], start=start):
        if el.tag == W + "tbl":
            ziel = i + 1
    vorrat = gruempi_pfade(GRUEMPI_RESERVE)

    for seite in range(anzahl):
        pfade = vorrat[seite * pro_seite:(seite + 1) * pro_seite]
        if not pfade:
            break

        def aufbau(pfade=pfade):
            par = doc.add_paragraph()
            par.paragraph_format.page_break_before = True
            bz._no_space(par)
            bz.add_photo_grid(doc, pfade, cols=2, max_h_cm=6.5)

        ziel = einsetzen(body, ziel, neue_elemente(doc, aufbau))


def _zeilenhoehe_cm(zeile):
    """Grobe Höhe einer Tabellenzeile in cm (Bilder exakt, Text geschätzt)."""
    hoch = 0.0
    for zelle in zeile.findall(W + "tc"):
        h = 0.0
        for par in zelle.findall(W + "p"):
            bilder = [Emu(int(e.get("cy"))).cm
                      for e in par.iter(WP + "extent")]
            h += max(bilder) if bilder else 0.5
        hoch = max(hoch, h)
    return hoch + 0.1


def tabellen_zusammenhalten(body, max_cm=20.0):
    """Tabellen nicht über den Seitenumbruch zerreissen.

    Jede Zeile bleibt in sich geschlossen (cantSplit). Tabellen, die auf eine
    Seite passen, werden zusätzlich als Block gehalten (keepNext auf allen
    Absätzen ausser denen der letzten Zeile) — reicht der Platz nicht, rückt
    die ganze Tabelle auf die nächste Seite. Fotoraster, die von Haus aus
    länger als eine Seite sind, bleiben umbrechbar.
    """
    ganz = geteilt = 0
    for tbl in body.iter(W + "tbl"):
        zeilen = tbl.findall(W + "tr")
        if not zeilen:
            continue
        for zeile in zeilen:
            trpr = zeile.find(W + "trPr")
            if trpr is None:
                trpr = etree.SubElement(zeile, W + "trPr")
                zeile.insert(0, trpr)
            if trpr.find(W + "cantSplit") is None:
                trpr.insert(0, zeile.makeelement(W + "cantSplit", {}))
        # Fotoraster fliessen weiter: die Bilder sollen direkt unter dem Text
        # beginnen, auch wenn nicht alle auf die Seite passen.
        bilder = tbl.find(".//" + W + "drawing") is not None
        if bilder and not absatztext(tbl).strip():
            geteilt += 1
            continue
        hoehe = sum(_zeilenhoehe_cm(z) for z in zeilen)
        if hoehe > max_cm:
            geteilt += 1
            continue
        for zeile in zeilen[:-1]:
            for par in zeile.iter(W + "p"):
                ppr = par.get_or_add_pPr()
                if ppr.find(W + "keepNext") is None:
                    ppr.insert(0, par.makeelement(W + "keepNext", {}))
        ganz += 1
    return ganz, geteilt


def sponsorenlogos(doc, zuordnung, hoehe_cm=1.9):
    """In jede Sponsorenbox die Logos aller genannten Sponsoren setzen.

    Die Logos kommen in ein Raster mit weissen Feldern (max. 3 pro Zeile), damit
    auch breite Wortmarken wie «MUOSER» gross genug bleiben. hoehe_cm ist die
    maximale Höhe eines Logos; breite Logos werden über die Feldbreite begrenzt.
    """
    gesetzt, fehlend = 0, []
    for tabelle in doc.tables:
        text = absatztext(tabelle._tbl)
        passend = next((v for k, v in zuordnung.items() if k in text), None)
        if passend is None:
            continue
        zelle = tabelle.rows[0].cells[0]
        # bisherige Logozeilen entfernen
        for par in zelle._tc.findall(W + "p"):
            if par.find(".//" + W + "drawing") is not None:
                entfernen(par)
        # Sponsorenzeilen aktualisieren / ergänzen
        for label, namen in passend.get("text", []):
            treffer = None
            for par in zelle.paragraphs:
                if absatztext(par._p).startswith(label + ":"):
                    treffer = par
                    break
            if treffer is not None:
                laeufe = treffer._p.findall(W + "r")
                for r in laeufe[1:]:
                    entfernen(r)
                lauf = treffer.add_run(namen)
                bz._set_font(lauf, size=9.5, color=bz.INK)
            else:
                neu = zelle.add_paragraph()
                bz._no_space(neu, after=1)
                lauf = neu.add_run(label + ": ")
                bz._set_font(lauf, size=9.5, bold=True, color=bz.INK)
                lauf = neu.add_run(namen)
                bz._set_font(lauf, size=9.5, color=bz.INK)
        pfade = []
        for name, datei in passend["logos"]:
            if datei is None:
                fehlend.append(name)
                continue
            pfad = os.path.join(REPO, "assets", "img", "sponsoren", datei)
            if os.path.exists(pfad):
                pfade.append(pfad)
            else:
                fehlend.append(name)
        if not pfade:
            continue
        # Logoraster: jedes Logo bekommt eine gleich grosse weisse Fläche und
        # wird darin so gross wie möglich abgebildet (Breite UND Höhe zählen).
        spalten = min(3, len(pfade))
        reihen = (len(pfade) + spalten - 1) // spalten
        raster = zelle.add_table(rows=reihen, cols=spalten)
        feld_b = 16.6 / spalten
        bz._table_fixed(raster, [Cm(feld_b)] * spalten)
        for i, pfad in enumerate(pfade):
            feld = raster.rows[i // spalten].cells[i % spalten]
            bz._shade_cell(feld, "FFFFFF")
            bz._cell_margins(feld, top=60, bottom=60, left=60, right=60)
            par = feld.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bz._no_space(par)
            bild = Image.open(pfad)
            breite = min(feld_b - 0.5, 6.0, hoehe_cm * bild.width / bild.height)
            par.add_run().add_picture(pfad, width=Cm(breite))
        for i in range(len(pfade), reihen * spalten):   # leere Felder weiss
            feld = raster.rows[i // spalten].cells[i % spalten]
            bz._shade_cell(feld, "FFFFFF")
        gesetzt += 1
    return gesetzt, fehlend


def agenda_tabulatoren(body, pos_cm=3.6):
    """Termine: Datum und Titel mit Tabulator trennen.

    Im Korrekturmodus wurden die Füllleerzeichen hinter dem Datum gelöscht —
    ohne Ersatz klebte «21.08.2026» direkt am Veranstaltungstitel.
    """
    start = body_index(body, lambda el: el.tag == W + "p" and absatztext(el)
                       .strip() == "Bevorstehende Veranstaltungen")
    if start < 0:
        return 0
    n = 0
    for par in list(body)[start + 1:start + 8]:
        if par.tag != W + "p":
            break
        runs = par.findall(W + "r")
        text = absatztext(par)
        if not text.strip():
            continue
        datum = DATUM.match(text)
        if not datum or len(runs) < 2:
            break
        # Tabulator hinter den letzten Run des Datums setzen (das Datum ist
        # durch die Korrekturen teils auf mehrere Runs verteilt)
        gesammelt = ""
        for i, run in enumerate(runs):
            gesammelt += "".join(t.text or "" for t in run.findall(W + "t"))
            if len(gesammelt.rstrip()) >= len(datum.group(1)):
                runs = runs[i:]
                break
        for t in runs[0].findall(W + "t"):
            if t.text and t.text.rstrip() != t.text:
                t.text = t.text.rstrip()
        # get_or_add_* hält die vom Schema verlangte Reihenfolge in pPr ein
        tabs = par.get_or_add_pPr().get_or_add_tabs()
        if not tabs.findall(W + "tab"):
            tab = tabs.add_tab()
            tab.set(W + "val", "left")
            tab.set(W + "pos", str(int(Cm(pos_cm).twips)))
        n += 1
        if par.findall(W + "r/" + W + "tab"):
            continue        # Tabulator ist da, es fehlte nur der Tabstopp
        lauf = par.makeelement(W + "r", {})
        rpr = runs[0].find(W + "rPr")
        if rpr is not None:
            lauf.append(copy.deepcopy(rpr))
        lauf.append(par.makeelement(W + "tab", {}))
        runs[0].addnext(lauf)
    return n


def toc_eintraege_entfernen(body, titel):
    """Einträge gelöschter Kapitel aus dem Inhaltsverzeichnis nehmen."""
    weg = 0
    for par in list(body.iter(W + "p")):
        stil = par.find(W + "pPr/" + W + "pStyle")
        if stil is None or not stil.get(W + "val").startswith(("Verzeichnis",
                                                              "TOC")):
            continue
        text = absatztext(par).strip()
        if any(text.startswith(t) for t in titel):
            entfernen(par)
            weg += 1
    return weg


def toc_seitenzahlen(body, seiten):
    """Zwischengespeicherte Seitenzahlen im Inhaltsverzeichnis setzen.

    seiten: {Eintragstext: Seitenzahl}. Nötig, weil ausser Word niemand die
    PAGEREF-Felder neu berechnet — im PDF stünden sonst alte Zahlen.
    """
    gesetzt = 0
    for par in body.iter(W + "p"):
        stil = par.find(W + "pPr/" + W + "pStyle")
        if stil is None or not stil.get(W + "val").startswith(("Verzeichnis",
                                                              "TOC")):
            continue
        texte = list(par.iter(W + "t"))
        if len(texte) < 2:
            continue
        titel = "".join(t.text or "" for t in texte[:-1]).strip()
        neu = seiten.get(titel.lower())
        if neu and texte[-1].text != str(neu):
            texte[-1].text = str(neu)
            gesetzt += 1
    return gesetzt


def bilder_optimieren(doc, dpi=220, qualitaet=88):
    """Eingebettete Bilder auf Druckauflösung bringen (Dateigrösse!).

    Ralph konnte die 40-MB-Fassung nicht mehr per Mail verschicken. Bilder
    werden deshalb auf die tatsächliche Anzeigegrösse × dpi verkleinert und
    Fotos ohne Transparenz als JPEG gespeichert.
    """
    # grösste Anzeigebreite je Bild (in EMU) aus dem Dokument holen
    breite = {}
    for blip in doc.element.body.iter(A + "blip"):
        rid = blip.get(R + "embed")
        rahmen = blip
        while rahmen is not None and rahmen.tag not in (WP + "inline",
                                                        WP + "anchor"):
            rahmen = rahmen.getparent()
        if rahmen is None:
            continue
        cx = int(rahmen.find(WP + "extent").get("cx"))
        breite[rid] = max(breite.get(rid, 0), cx)

    vorher = nachher = 0
    for rid, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue
        teil = rel.target_part
        alt = teil.blob
        vorher += len(alt)
        try:
            bild = Image.open(io.BytesIO(alt))
        except Exception:
            nachher += len(alt)
            continue
        cm = Emu(breite.get(rid, 0)).cm or 5.0
        transparent = (bild.mode in ("RGBA", "LA", "P")
                       and "transparency" not in bild.info
                       and bild.mode != "P"
                       and bild.getchannel("A").getextrema()[0] < 250)
        if bild.mode == "P":
            transparent = "transparency" in bild.info
        # Logos/Strichzeichnungen brauchen mehr Auflösung als Fotos
        ziel_px = max(int(cm / 2.54 * (350 if transparent else dpi)), 400)
        if bild.width > ziel_px:
            hoehe = int(round(bild.height * ziel_px / bild.width))
            bild = bild.resize((ziel_px, hoehe), Image.LANCZOS)
        if transparent:
            neu = _als_bytes(bild.convert("RGBA"), "PNG")
            if len(neu) < len(alt):
                teil._blob = neu
        else:
            if bild.mode in ("RGBA", "LA", "P"):
                flaeche = Image.new("RGB", bild.size, "white")
                bild = bild.convert("RGBA")
                flaeche.paste(bild, mask=bild.split()[3])
                bild = flaeche
            neu = _als_bytes(bild, "JPEG")
            if len(neu) < len(alt):
                teil._blob = neu
                if teil.partname.ext.lower() != "jpg":
                    teil.partname = docx.opc.packuri.PackURI(
                        "/word/media/opt%s.jpg" % rid)
                    teil._content_type = "image/jpeg"
        nachher += len(teil.blob)
    print("10) Bilder optimiert: %.1f MB → %.1f MB (max %d dpi)"
          % (vorher / 1e6, nachher / 1e6, dpi))


def schattierung_robust(doc):
    """Bei jeder Flächenfarbe die Musterfarbe auf die Füllfarbe setzen.

    Word wertet bei w:val="clear" nur w:fill aus. Schwächere Renderer (Vorschau,
    Pages, Web-Viewer) mischen dagegen die Musterfarbe «auto» ein — daher die
    ausgewaschenen Grau-/Lachstöne. Mit w:color = w:fill stimmt die Farbe in
    beiden Fällen.
    """
    n = 0
    for teil in doc.part.package.iter_parts():
        wurzel = getattr(teil, "element", None)
        if wurzel is None:
            continue
        for shd in wurzel.iter(W + "shd"):
            fuellung = shd.get(W + "fill")
            if not fuellung or fuellung in ("auto",):
                continue
            if shd.get(W + "color") in (None, "auto"):
                shd.set(W + "color", fuellung)
                n += 1
    return n


# Kopfzeile: kräftigere Vereinsfarben (Wunsch Joel, 28.07.2026)
KOPF_SCHWARZ = "000000"      # statt 181818
KOPF_ROT = "C1121F"          # blutrot statt E63124


def kopfzeilen_farben(doc, schwarz=KOPF_SCHWARZ, rot=KOPF_ROT):
    """Balken und Trennlinie der Kopfzeile in kräftigeren Vereinsfarben."""
    ersatz = {"181818": schwarz, "E63124": rot, "D9261C": rot}
    n = 0
    for teil in doc.part.package.iter_parts():
        if "header" not in str(teil.partname) and "footer" not in str(teil.partname):
            continue
        wurzel = getattr(teil, "element", None)
        if wurzel is None:
            continue
        for shd in wurzel.iter(W + "shd"):
            neu = ersatz.get((shd.get(W + "fill") or "").upper())
            if neu:
                shd.set(W + "fill", neu)
                shd.set(W + "color", neu)
                n += 1
    return n


def datei_offen(pfad):
    """Hat ein Programm (Word …) die Datei gerade offen?

    Wird die .docx überschrieben, während Word sie geöffnet hat, lädt Word die
    Bilder aus dem alten Paket nicht mehr — im Dokument stehen dann leere
    Rahmen mit «Das Bild kann nicht angezeigt werden».
    """
    try:
        return bool(subprocess.check_output(
            ["lsof", "--", pfad], stderr=subprocess.DEVNULL).strip())
    except (subprocess.CalledProcessError, OSError):
        return False


def verwaiste_bilder_entfernen(doc):
    """Nicht mehr referenzierte Bildteile aus dem Paket werfen (Dateigrösse)."""
    benutzt = {b.get(R + "embed") for b in doc.element.body.iter(A + "blip")}
    benutzt |= {b.get(R + "link") for b in doc.element.body.iter(A + "blip")}
    weg = 0
    for rid, rel in list(doc.part.rels.items()):
        if "image" in rel.reltype and rid not in benutzt:
            doc.part.drop_rel(rid)
            weg += 1
    return weg


def felder_aktualisieren(doc):
    """Word beim Öffnen zum Aktualisieren des Inhaltsverzeichnisses auffordern."""
    settings = doc.settings.element
    if settings.find(W + "updateFields") is None:
        el = etree.SubElement(settings, W + "updateFields")
        el.set(W + "val", "true")


# ---------------------------------------------------------------------- Ablauf


def main(fueller=0, ziel=ZIEL, seiten=None):
    doc = docx.Document(QUELLE)
    body = doc.element.body

    stat = accept_revisions(doc.element)
    print("1) Korrekturmodus angenommen:", stat)

    # ---- Titelseite: pixeliges Vereinslogo (96 × 75 px) ersetzen
    logo_alt = os.path.join(REPO, "output", "img-cache", "cover-logo-alt.png")
    rid = teil = None
    for rid_, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue
        blob = rel.target_part.blob
        try:
            bild = Image.open(io.BytesIO(blob))
        except Exception:
            continue
        if bild.size == (96, 75):
            rid, teil = rid_, rel.target_part
            break
    if teil is None:
        print("2) WARNUNG: Titelseiten-Logo nicht gefunden")
    else:
        os.makedirs(os.path.dirname(logo_alt), exist_ok=True)
        open(logo_alt, "wb").write(teil.blob)
        scharf = logo_rendern(
            os.path.join(REPO, "assets", "img", "fcs-logo-wappen.svg"))
        bild_ersetzen(doc, rid, teil, scharf, hoehe_behalten=True)
        print("2) Titelseiten-Logo ersetzt: 96 × 75 px → %d × %d px"
              % scharf.size)

    # ---- 3. Mannschaft: neues Teamfoto + neuer Dress-Sponsor
    rid, teil = bildteil_suchen(doc, os.path.join(
        REPO, "output", "img-cache", "FCS_3_jpg.jpg"))
    if teil is None:
        print("4) WARNUNG: Teamfoto FCS 3 nicht gefunden")
    else:
        neu = ImageOps.exif_transpose(
            Image.open(os.path.join(SRC, "Fotos", "3 Mannschaft.png")))
        neu.thumbnail((1600, 1600), Image.LANCZOS)
        bild_ersetzen(doc, rid, teil, neu)
        print("4) Teamfoto 3. Mannschaft ersetzt (%d × %d px)" % neu.size)
    print("   Dress-Sponsor angepasst:",
          text_ersetzen(body, "BINARY one GmbH", "Feritec AG"), "Stelle(n)")

    # ---- Redaktionelle Hinweise entfernen
    weg = absaetze_entfernen(body, "[PRÜFEN]")
    weg += absaetze_entfernen(body, "[PLATZHALTER]")
    runs = run_entfernen(body, "[PRÜFEN — Stand letzte Ausgabe]")
    runs += run_entfernen(body, "[PLATZHALTER — durch Sponsoringverantwortlichen")
    for par in list(body.iter(W + "p")):
        if absatztext(par).strip() == "Weitere Sponsoren Juniorenabteilung:":
            entfernen(par)
    print("5) Redaktionshinweise entfernt: %d Absätze, %d Textstellen"
          % (weg, runs))

    # ---- Leere Tabellenreste (z. B. gelöschte FCS-3-Statistik)
    leer = 0
    for tbl in list(body.iter(W + "tbl")):
        if not tbl.findall(W + "tr"):
            entfernen(tbl)
            leer += 1
    print("6) Leere Tabellenreste entfernt:", leer)

    # ---- Sponsorenlogos in die Sponsorenboxen
    gesetzt, fehlend = sponsorenlogos(doc, SPONSOREN)
    print("7) Sponsorenboxen mit Logos bestückt:", gesetzt,
          "| ohne Logo:", ", ".join(fehlend) if fehlend else "—")

    # ---- Grümpelturnier-Bilder
    events_in_bildern(doc)
    print("   «Events in Bildern» mit %d Grümpi-Fotos bestückt"
          % (len(GRUEMPI_QUER) + len(GRUEMPI_HOCH)))

    print("   Termine mit Tabulator ausgerichtet:",
          agenda_tabulatoren(body), "Zeilen")

    # ---- Inhaltsverzeichnis: gelöschte Kapitel raus
    print("8) Inhaltsverzeichnis bereinigt:",
          toc_eintraege_entfernen(body, ("Sponsoring", "Administration")),
          "Einträge")

    # ---- Seitenausgleich mit ganzen Fotoseiten
    fueller_seiten(doc, fueller)
    if fueller:
        print("   %d Fotoseite(n) für die Seitenzahl ergänzt" % fueller)

    print("   verwaiste Bilder entfernt:", verwaiste_bilder_entfernen(doc))
    print("   Flächenfarben abgesichert (Musterfarbe = Füllfarbe):",
          schattierung_robust(doc))
    print("   Kopfzeile auf %s / %s gesetzt:" % (KOPF_SCHWARZ, KOPF_ROT),
          kopfzeilen_farben(doc), "Flächen")

    ganz, fliessend = tabellen_zusammenhalten(body)
    print("9) Tabellen zusammengehalten:", ganz,
          "| Fotoraster/zu lang → fliessend:", fliessend)

    bilder_optimieren(doc)

    if seiten:
        with open(seiten, encoding="utf-8") as fh:
            zahlen = {k.lower(): v for k, v in json.load(fh).items()}
        print("   Seitenzahlen im Inhaltsverzeichnis gesetzt:",
              toc_seitenzahlen(body, zahlen))

    felder_aktualisieren(doc)
    if datei_offen(ziel):
        ziel = os.path.splitext(ziel)[0] + " (neu).docx"
        print("!  Zieldatei ist in Word geöffnet — geschrieben wird stattdessen:")
    doc.save(ziel)
    print("Fertig:", ziel)


if __name__ == "__main__":
    n = 0
    if "--fueller" in sys.argv:
        n = int(sys.argv[sys.argv.index("--fueller") + 1])
    zielpfad = ZIEL
    if "--ziel" in sys.argv:
        zielpfad = sys.argv[sys.argv.index("--ziel") + 1]
    seitendatei = None
    if "--seiten" in sys.argv:
        seitendatei = sys.argv[sys.argv.index("--seiten") + 1]
    sys.exit(main(n, zielpfad, seitendatei))
